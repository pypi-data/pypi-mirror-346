#!/usr/bin/env python3
"""
Enhanced Citation Manager

A comprehensive command-line tool for managing citations from various sources
and in multiple formats.

Dependencies:
- requests
- python-docx (for Word document processing)
- beautifulsoup4 (for URL parsing)
"""

import requests
import json
import sys
import os
import re
import argparse
from typing import Dict, Any, List, Optional, Tuple
import xml.etree.ElementTree as ET
from datetime import datetime
from urllib.parse import urlparse
from docx import Document


class BaseCitationFormatter:
    """Base class for citation formatters"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a citation string"""
        raise NotImplementedError("Subclasses must implement this method")
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation"""
        raise NotImplementedError("Subclasses must implement this method")
    
class ElsevierFormatter(BaseCitationFormatter):
    """Format citations in Elsevier style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an Elsevier-style citation"""
        # Elsevier typically uses the Vancouver system
        
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 6:
                authors = ", ".join(metadata["authors"][:6]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
        
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')};"
        
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')})"
            
        if metadata.get("pages"):
            citation += f":{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Elsevier"""
        if metadata.get("authors"):
            first_author = metadata["authors"][0].split(",")[0]
            if len(metadata["authors"]) > 1:
                return f"({first_author} et al., {metadata.get('year', '')})"
            else:
                return f"({first_author}, {metadata.get('year', '')})"
        else:
            return f"(Anonymous, {metadata.get('year', '')})"

class SpringerFormatter(BaseCitationFormatter):
    """Format citations in Springer style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Springer-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 3:
                authors = metadata["authors"][0] + " et al"
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}: {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}"
            
        if metadata.get("volume"):
            citation += f" {metadata.get('volume', '')}"
            
        if metadata.get("pages"):
            citation += f", {metadata.get('pages', '')}"
            
        if metadata.get("year"):
            citation += f" ({metadata.get('year', '')})"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Springer"""
        if metadata.get("authors"):
            first_author = metadata["authors"][0].split(",")[0]
            if len(metadata["authors"]) > 1:
                return f"{first_author} et al. ({metadata.get('year', '')})"
            else:
                return f"{first_author} ({metadata.get('year', '')})"
        else:
            return f"Anonymous ({metadata.get('year', '')})"

class APAFormatter(BaseCitationFormatter):
    """Format citations in APA style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an APA-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 7:
                authors = ", ".join(metadata["authors"][:6]) + ", ... " + metadata["authors"][-1]
            else:
                authors = ", ".join(metadata["authors"])
            
            # Replace last comma with "&"
            last_comma_index = authors.rfind(",")
            if last_comma_index != -1:
                authors = authors[:last_comma_index] + " &" + authors[last_comma_index+1:]
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. ({metadata.get('year', '')}). {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}"
            
            if metadata.get("volume"):
                citation += f", {metadata.get('volume', '')}"
                
                if metadata.get("issue"):
                    citation += f"({metadata.get('issue', '')})"
                    
            if metadata.get("pages"):
                citation += f", {metadata.get('pages', '')}"
        
        citation += "."
        
        if metadata.get("doi"):
            citation += f" https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for APA"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"({author}, {metadata.get('year', '')})"
            elif len(metadata["authors"]) == 2:
                author1 = metadata["authors"][0].split(",")[0]
                author2 = metadata["authors"][1].split(",")[0]
                return f"({author1} & {author2}, {metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"({author} et al., {metadata.get('year', '')})"
        else:
            return f"(Anonymous, {metadata.get('year', '')})"

# Additional citation formatters for various journal styles

class NatureFormatter(BaseCitationFormatter):
    """Format citations in Nature style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Nature-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 5:
                authors = ", ".join(metadata["authors"][:5]) + " et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors} "
        citation += f"{metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')} "
            
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')}). "
        
        if metadata.get("doi"):
            citation += f"https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Nature"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} ({metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al. ({metadata.get('year', '')})"
        else:
            return f"Anonymous ({metadata.get('year', '')})"

class ScienceFormatter(BaseCitationFormatter):
    """Format citations in Science style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Science-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 5:
                authors = ", ".join(metadata["authors"][:5]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}, "
        citation += f"{metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}. "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')} "
            
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')}). "
        
        if metadata.get("doi"):
            citation += f"doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Science"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} ({metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al. ({metadata.get('year', '')})"
        else:
            return f"Anonymous ({metadata.get('year', '')})"

class IEEEFormatter(BaseCitationFormatter):
    """Format citations in IEEE style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an IEEE-style citation"""
        # Format authors with initials
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = "".join([name[0] + "." for name in first_name.split()])
                    formatted_authors.append(f"{last_name}, {initials}")
                else:
                    formatted_authors.append(author)
            
            if len(formatted_authors) > 3:
                authors = ", ".join(formatted_authors[:3]) + ", et al."
            else:
                authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}, \"{metadata.get('title', '')},\" "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}, "
            
        if metadata.get("volume"):
            citation += f"vol. {metadata.get('volume', '')}, "
            
        if metadata.get("issue"):
            citation += f"no. {metadata.get('issue', '')}, "
            
        if metadata.get("pages"):
            citation += f"pp. {metadata.get('pages', '').replace('-', '–')}, "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}. "
        
        if metadata.get("doi"):
            citation += f"doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for IEEE"""
        # IEEE uses numbered citations, but we'll return author-date for consistency
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"[{author}, {metadata.get('year', '')}]"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"[{author} et al., {metadata.get('year', '')}]"
        else:
            return f"[Anonymous, {metadata.get('year', '')}]"

class ChicagoFormatter(BaseCitationFormatter):
    """Format citations in Chicago style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Chicago-style citation"""
        # Format authors
        if metadata.get("authors"):
            # Flip first author only
            formatted_authors = []
            for i, author in enumerate(metadata["authors"]):
                parts = author.split(",")
                if i == 0 and len(parts) == 2:
                    # Keep first author as Last, First
                    formatted_authors.append(author)
                elif len(parts) == 2:
                    # Flip subsequent authors to First Last
                    formatted_authors.append(f"{parts[1].strip()} {parts[0].strip()}")
                else:
                    formatted_authors.append(author)
            
            if len(formatted_authors) > 10:
                authors = ", ".join(formatted_authors[:7]) + ", et al."
            else:
                authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. \"{metadata.get('title', '')}.\" "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("issue"):
            citation += f"no. {metadata.get('issue', '')} "
            
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')}): "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')}. "
        else:
            citation += ". "
        
        if metadata.get("doi"):
            citation += f"https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Chicago"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"({author} {metadata.get('year', '')})"
            elif len(metadata["authors"]) == 2:
                author1 = metadata["authors"][0].split(",")[0]
                author2 = metadata["authors"][1].split(",")[0]
                return f"({author1} and {author2} {metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"({author} et al. {metadata.get('year', '')})"
        else:
            return f"(Anonymous {metadata.get('year', '')})"

class MLAFormatter(BaseCitationFormatter):
    """Format citations in MLA style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an MLA-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                authors = metadata["authors"][0]
            elif len(metadata["authors"]) == 2:
                parts1 = metadata["authors"][0].split(",")
                parts2 = metadata["authors"][1].split(",")
                
                if len(parts1) == 2 and len(parts2) == 2:
                    # First author as Last, First; second as First Last
                    authors = f"{parts1[0]}, {parts1[1].strip()} and {parts2[1].strip()} {parts2[0]}"
                else:
                    authors = f"{metadata['authors'][0]} and {metadata['authors'][1]}"
            elif len(metadata["authors"]) > 2:
                authors = f"{metadata['authors'][0]}, et al."
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. \"{metadata.get('title', '')}.\" "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}, "
            
        if metadata.get("volume"):
            citation += f"vol. {metadata.get('volume', '')}, "
            
        if metadata.get("issue"):
            citation += f"no. {metadata.get('issue', '')}, "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}, "
            
        if metadata.get("pages"):
            citation += f"pp. {metadata.get('pages', '')}. "
        else:
            citation += ". "
        
        if metadata.get("doi"):
            citation += f"DOI: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for MLA"""
        if metadata.get("authors"):
            author = metadata["authors"][0].split(",")[0]
            return f"({author})"
        else:
            return "(Anonymous)"

class HarvardFormatter(BaseCitationFormatter):
    """Format citations in Harvard style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Harvard-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = "".join([name[0] + "." for name in first_name.split()])
                    formatted_authors.append(f"{last_name}, {initials}")
                else:
                    formatted_authors.append(author)
            
            if len(formatted_authors) > 3:
                authors = ", ".join(formatted_authors[:3]) + ", et al."
            else:
                authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors} "
        
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')}). "
            
        citation += f"{metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}, "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')}), "
        else:
            citation += ", "
            
        if metadata.get("pages"):
            citation += f"pp. {metadata.get('pages', '')}. "
        else:
            citation += ". "
        
        if metadata.get("doi"):
            citation += f"doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Harvard"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"({author}, {metadata.get('year', '')})"
            elif len(metadata["authors"]) == 2:
                author1 = metadata["authors"][0].split(",")[0]
                author2 = metadata["authors"][1].split(",")[0]
                return f"({author1} and {author2}, {metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"({author} et al., {metadata.get('year', '')})"
        else:
            return f"(Anonymous, {metadata.get('year', '')})"

class VancouverFormatter(BaseCitationFormatter):
    """Format citations in Vancouver style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Vancouver-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = "".join([name[0] for name in first_name.split()])
                    formatted_authors.append(f"{last_name} {initials}")
                else:
                    formatted_authors.append(author)
            
            if len(formatted_authors) > 6:
                authors = " ".join(formatted_authors[:6]) + ", et al."
            else:
                authors = " ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}. "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}"
            
        if metadata.get("volume"):
            citation += f";{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')})"
            
        if metadata.get("pages"):
            citation += f":{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Vancouver (numbered)"""
        # Vancouver typically uses a numbered reference system
        # This is a placeholder since we don't track reference numbers
        return "(N)"

class BMCFormatter(BaseCitationFormatter):
    """Format citations in BMC style (BioMed Central)"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a BMC-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 30:
                authors = ", ".join(metadata["authors"][:30]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}: {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}"
            
        if metadata.get("volume"):
            citation += f", {metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')})"
            
        if metadata.get("pages"):
            citation += f":{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for BMC"""
        if metadata.get("authors"):
            first_author = metadata["authors"][0].split(",")[0]
            if len(metadata["authors"]) > 1:
                return f"{first_author} et al. [{metadata.get('year', '')}]"
            else:
                return f"{first_author} [{metadata.get('year', '')}]"
        else:
            return f"Anonymous [{metadata.get('year', '')}]"

class PLOSFormatter(BaseCitationFormatter):
    """Format citations in PLOS style (Public Library of Science)"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a PLOS-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 10:
                authors = ", ".join(metadata["authors"][:10]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors} "
        
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')}) "
            
        citation += f"{metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')}): "
        else:
            citation += ": "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')}. "
        else:
            citation += ". "
        
        if metadata.get("doi"):
            citation += f"https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for PLOS"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author}, {metadata.get('year', '')}"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al., {metadata.get('year', '')}"
        else:
            return f"Anonymous, {metadata.get('year', '')}"

class CellFormatter(BaseCitationFormatter):
    """Format citations in Cell style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a Cell-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = "".join([name[0] + "." for name in first_name.split()])
                    formatted_authors.append(f"{last_name}, {initials}")
                else:
                    formatted_authors.append(author)
            
            authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors} "
        
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')}). "
            
        citation += f"{metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')}. "
        else:
            citation += ". "
        
        if metadata.get("doi"):
            citation += f"https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Cell"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"({author}, {metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"({author} et al., {metadata.get('year', '')})"
        else:
            return f"(Anonymous, {metadata.get('year', '')})"
        
class JAMAFormatter(BaseCitationFormatter):
    """Format citations in JAMA (Journal of the American Medical Association) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a JAMA-style citation"""
        # Format authors - JAMA uses et al after 6 authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 6:
                authors = ", ".join(metadata["authors"][:6]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}. "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}"
            
        if metadata.get("volume"):
            citation += f";{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')})"
            
        if metadata.get("pages"):
            citation += f":{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi:{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for JAMA"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author}"
            elif len(metadata["authors"]) == 2:
                author1 = metadata["authors"][0].split(",")[0]
                author2 = metadata["authors"][1].split(",")[0]
                return f"{author1} and {author2}"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al"
        else:
            return "Anonymous"

class BMJFormatter(BaseCitationFormatter):
    """Format citations in BMJ (British Medical Journal) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a BMJ-style citation"""
        # Format authors - BMJ uses et al after 6 authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 6:
                authors = ", ".join(metadata["authors"][:6]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}"
            
        if metadata.get("volume"):
            citation += f";{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')})"
            
        if metadata.get("pages"):
            citation += f":{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for BMJ"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author}"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al"
        else:
            return "Anonymous"

class NEJMFormatter(BaseCitationFormatter):
    """Format citations in NEJM (New England Journal of Medicine) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a NEJM-style citation"""
        # Format authors - NEJM lists all authors
        if metadata.get("authors"):
            authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}"
            
        if metadata.get("volume"):
            citation += f";{metadata.get('volume', '')}"
            
        if metadata.get("pages"):
            citation += f":{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" DOI: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for NEJM"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al."
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al."
        else:
            return "Anonymous"

class JBCFormatter(BaseCitationFormatter):
    """Format citations in JBC (Journal of Biological Chemistry) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into a JBC-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = "".join([name[0] + "." for name in first_name.split()])
                    formatted_authors.append(f"{last_name}, {initials}")
                else:
                    formatted_authors.append(author)
            
            authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors} ({metadata.get('year', '')}) {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for JBC"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} ({metadata.get('year', '')})"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al. ({metadata.get('year', '')})"
        else:
            return f"Anonymous ({metadata.get('year', '')})"

class ACMFormatter(BaseCitationFormatter):
    """Format citations in ACM (Association for Computing Machinery) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an ACM-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = " ".join([name[0] + "." for name in first_name.split()])
                    formatted_authors.append(f"{first_name[0]}. {last_name}")
                else:
                    formatted_authors.append(author)
            
            authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}. {metadata.get('year', '')}. {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("issue"):
            citation += f"{metadata.get('issue', '')} "
            
        if metadata.get("pages"):
            citation += f"({metadata.get('year', '')}), {metadata.get('pages', '')}. "
        else:
            citation += f"({metadata.get('year', '')}). "
        
        if metadata.get("doi"):
            citation += f"https://doi.org/{metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for ACM"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"[{author} {metadata.get('year', '')}]"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"[{author} et al. {metadata.get('year', '')}]"
        else:
            return f"[Anonymous {metadata.get('year', '')}]"

class OxfordFormatter(BaseCitationFormatter):
    """Format citations in Oxford style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into Oxford-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    formatted_authors.append(f"{last_name}, {first_name[0]}.")
                else:
                    formatted_authors.append(author)
            
            if len(formatted_authors) > 4:
                authors = ", ".join(formatted_authors[:4]) + ", et al."
            else:
                authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}, '{metadata.get('title', '')}', "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}, "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"/{metadata.get('issue', '')}"
            
        citation += f" ({metadata.get('year', '')}), "
            
        if metadata.get("pages"):
            citation += f"pp. {metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for Oxford"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author}, {metadata.get('year', '')}"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al., {metadata.get('year', '')}"
        else:
            return f"Anonymous, {metadata.get('year', '')}"

class RSCFormatter(BaseCitationFormatter):
    """Format citations in RSC (Royal Society of Chemistry) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an RSC-style citation"""
        # Format authors
        if metadata.get("authors"):
            if len(metadata["authors"]) > 30:
                authors = ", ".join(metadata["authors"][:30]) + ", et al."
            else:
                authors = ", ".join(metadata["authors"])
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}, {metadata.get('title', '')}, "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')}, "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}, "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" DOI: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for RSC"""
        return f"{metadata.get('journal', '')} {metadata.get('year', '')}"

class ACSFormatter(BaseCitationFormatter):
    """Format citations in ACS (American Chemical Society) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an ACS-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials
                    initials = "".join([name[0] + "." for name in first_name.split()]) 
                    formatted_authors.append(f"{last_name}, {initials}")
                else:
                    formatted_authors.append(author)
            
            authors = "; ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors} {metadata.get('title', '')}. "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("year"):
            citation += f"{metadata.get('year', '')}, "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}"
            
        if metadata.get("issue"):
            citation += f"({metadata.get('issue', '')}), "
        else:
            citation += ", "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')}"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" DOI: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for ACS"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author}, {metadata.get('year', '')}"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al., {metadata.get('year', '')}"
        else:
            return f"Anonymous, {metadata.get('year', '')}"

class AIPFormatter(BaseCitationFormatter):
    """Format citations in AIP (American Institute of Physics) style"""
    
    def format_citation(self, metadata: Dict[str, Any]) -> str:
        """Format metadata into an AIP-style citation"""
        # Format authors
        if metadata.get("authors"):
            formatted_authors = []
            for author in metadata["authors"]:
                parts = author.split(",")
                if len(parts) == 2:
                    last_name = parts[0].strip()
                    first_name = parts[1].strip()
                    # Get initials with spaces
                    initials = " ".join([name[0] + "." for name in first_name.split()]) 
                    formatted_authors.append(f"{first_name[0]}. {last_name}")
                else:
                    formatted_authors.append(author)
            
            authors = ", ".join(formatted_authors)
        else:
            authors = "Anonymous"
        
        # Build the citation
        citation = f"{authors}, \"{metadata.get('title', '')},\" "
        
        if metadata.get("journal"):
            citation += f"{metadata.get('journal', '')} "
            
        if metadata.get("volume"):
            citation += f"{metadata.get('volume', '')}, "
            
        if metadata.get("pages"):
            citation += f"{metadata.get('pages', '')} "
            
        if metadata.get("year"):
            citation += f"({metadata.get('year', '')})"
            
        citation += "."
        
        if metadata.get("doi"):
            citation += f" doi: {metadata.get('doi', '')}"
        
        return citation
    
    def format_in_text_citation(self, metadata: Dict[str, Any]) -> str:
        """Format in-text citation for AIP"""
        if metadata.get("authors"):
            if len(metadata["authors"]) == 1:
                author = metadata["authors"][0].split(",")[0]
                return f"{author}, {metadata.get('year', '')}"
            else:
                author = metadata["authors"][0].split(",")[0]
                return f"{author} et al., {metadata.get('year', '')}"
        else:
            return f"Anonymous, {metadata.get('year', '')}"

class EnhancedMetadataExtractor:
    """Enhanced metadata extractor with support for more sources"""
    
    def __init__(self):
        self.api_keys = {
            "scopus": "",  # Add your Scopus API key here if available
            "semantic_scholar": "",  # Add your Semantic Scholar API key here if available
        }
    
    def extract_from_doi(self, doi: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a DOI using the Crossref API"""
        url = f"https://api.crossref.org/works/{doi}"
        headers = {"Accept": "application/json"}
        
        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()
            
            if "message" in data:
                return self._parse_crossref_data(data["message"])
            
            # If Crossref fails, try DataCite
            return self._try_datacite_doi(doi)
        except Exception as e:
            print(f"Error extracting metadata from Crossref DOI {doi}: {e}")
            # Try DataCite as a fallback
            return self._try_datacite_doi(doi)
    
    def _try_datacite_doi(self, doi: str) -> Optional[Dict[str, Any]]:
        """Try to get metadata from DataCite API"""
        url = f"https://api.datacite.org/dois/{doi}"
        
        try:
            response = requests.get(url)
            response.raise_for_status()
            data = response.json()
            
            if "data" in data and "attributes" in data["data"]:
                return self._parse_datacite_data(data["data"]["attributes"])
            return None
        except Exception as e:
            print(f"Error extracting metadata from DataCite DOI {doi}: {e}")
            return None
    
    def extract_from_pmid(self, pmid: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a PMID using the NCBI E-utilities API"""
        base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
        params = {
            "db": "pubmed",
            "id": pmid,
            "retmode": "json"
        }
        
        try:
            response = requests.get(base_url, params=params)
            response.raise_for_status()
            data = response.json()
            
            if "result" in data and pmid in data["result"]:
                return self._parse_pubmed_data(data["result"][pmid])
            return None
        except Exception as e:
            print(f"Error extracting metadata from PMID {pmid}: {e}")
            return None
    
    def extract_from_arxiv(self, arxiv_id: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from an arXiv ID using the arXiv API"""
        # Clean the arXiv ID (remove version if present)
        if "v" in arxiv_id:
            arxiv_id = arxiv_id.split("v")[0]
        
        url = f"http://export.arxiv.org/api/query?id_list={arxiv_id}"
        
        try:
            response = requests.get(url)
            response.raise_for_status()
            
            # Parse the XML response
            root = ET.fromstring(response.content)
            
            # Find the entry element
            ns = {"atom": "http://www.w3.org/2005/Atom"}
            entry = root.find(".//atom:entry", ns)
            
            if entry is not None:
                return self._parse_arxiv_entry(entry, ns)
            return None
        except Exception as e:
            print(f"Error extracting metadata from arXiv ID {arxiv_id}: {e}")
            return None
    
    def extract_from_isbn(self, isbn: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from an ISBN using the Open Library API"""
        url = f"https://openlibrary.org/api/books?bibkeys=ISBN:{isbn}&format=json&jscmd=data"
        
        try:
            response = requests.get(url)
            response.raise_for_status()
            data = response.json()
            
            key = f"ISBN:{isbn}"
            if key in data:
                return self._parse_openlibrary_data(data[key])
            return None
        except Exception as e:
            print(f"Error extracting metadata from ISBN {isbn}: {e}")
            return None
    
    # Update the extract_from_url method in the EnhancedMetadataExtractor class

    def extract_from_url(self, url: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a URL"""
        # Check if it's a PMC URL
        if "pmc.ncbi.nlm.nih.gov/articles/PMC" in url:
            # Extract PMC ID
            pmc_match = re.search(r'PMC(\d+)', url)
            if pmc_match:
                pmc_id = pmc_match.group(1)
                # Try to convert PMC ID to PMID using NCBI's eutils
                try:
                    convert_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?db=pubmed&term=PMC{pmc_id}[pmcid]&retmode=json"
                    response = requests.get(
                        convert_url,
                        headers={
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
                        }
                    )
                    data = response.json()
                    if "esearchresult" in data and "idlist" in data["esearchresult"] and data["esearchresult"]["idlist"]:
                        pmid = data["esearchresult"]["idlist"][0]
                        return self.extract_from_pmid(pmid)
                except Exception as e:
                    print(f"Error converting PMC{pmc_id} to PMID: {e}")
                    # Try direct extraction with proper headers
                    return self._extract_from_generic_url_with_headers(url)
        
        # Check if it's a known publisher URL
        if "nature.com" in url:
            return self._extract_from_nature_url(url)
        elif "science.org" in url:
            return self._extract_from_science_url(url)
        elif "pubmed.ncbi.nlm.nih.gov" in url:
            # Extract PMID from PubMed URL
            pmid = url.split("/")[-1].split("?")[0]
            return self.extract_from_pmid(pmid)
        elif "doi.org" in url:
            # Extract DOI from DOI URL
            doi = url.split("doi.org/")[-1]
            return self.extract_from_doi(doi)
        elif "arxiv.org" in url:
            # Extract arXiv ID from arXiv URL
            if "abs" in url:
                arxiv_id = url.split("abs/")[-1].split("v")[0]
                return self.extract_from_arxiv(arxiv_id)
        
        # Generic URL handling - use web scraping with meta tags
        return self._extract_from_generic_url_with_headers(url)

    def _extract_from_generic_url_with_headers(self, url: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a generic URL using meta tags with proper headers"""
        try:
            # Use a more browser-like User-Agent
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.5"
            }
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            # Use BeautifulSoup to parse HTML if available
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.text, 'html.parser')
                
                metadata = {
                    "title": "",
                    "authors": [],
                    "journal": "",
                    "year": "",
                    "url": url
                }
                
                # Extract title
                title_tag = soup.find("meta", property="og:title") or soup.find("meta", attrs={"name": "citation_title"})
                if title_tag:
                    metadata["title"] = title_tag.get("content", "")
                else:
                    title_tag = soup.find("title")
                    if title_tag:
                        metadata["title"] = title_tag.text
                
                # Extract authors
                author_tags = soup.find_all("meta", attrs={"name": "citation_author"})
                if author_tags:
                    metadata["authors"] = [tag.get("content", "") for tag in author_tags]
                
                # Extract journal
                journal_tag = soup.find("meta", attrs={"name": "citation_journal_title"})
                if journal_tag:
                    metadata["journal"] = journal_tag.get("content", "")
                
                # Extract year
                date_tag = soup.find("meta", attrs={"name": "citation_publication_date"})
                if date_tag:
                    date_content = date_tag.get("content", "")
                    if date_content and len(date_content) >= 4:
                        metadata["year"] = date_content[:4]
                
                # Extract DOI
                doi_tag = soup.find("meta", attrs={"name": "citation_doi"})
                if doi_tag:
                    metadata["doi"] = doi_tag.get("content", "")
                
                return metadata
            except ImportError:
                # If BeautifulSoup is not available, use a simpler approach
                print("BeautifulSoup not available, using simple extraction")
                metadata = {
                    "title": "",
                    "authors": [],
                    "url": url
                }
                
                # Simple title extraction
                title_match = re.search(r"<title>(.*?)</title>", response.text)
                if title_match:
                    metadata["title"] = title_match.group(1)
                
                return metadata
        except Exception as e:
            print(f"Error extracting metadata from URL {url}: {e}")
            return None
    
    def _extract_from_generic_url(self, url: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a generic URL using meta tags"""
        try:
            response = requests.get(url)
            response.raise_for_status()
            
            # Use BeautifulSoup to parse HTML if available
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.text, 'html.parser')
                
                metadata = {
                    "title": "",
                    "authors": [],
                    "journal": "",
                    "year": "",
                    "url": url
                }
                
                # Extract title
                title_tag = soup.find("meta", property="og:title") or soup.find("meta", attrs={"name": "citation_title"})
                if title_tag:
                    metadata["title"] = title_tag.get("content", "")
                else:
                    title_tag = soup.find("title")
                    if title_tag:
                        metadata["title"] = title_tag.text
                
                # Extract authors
                author_tags = soup.find_all("meta", attrs={"name": "citation_author"})
                if author_tags:
                    metadata["authors"] = [tag.get("content", "") for tag in author_tags]
                
                # Extract journal
                journal_tag = soup.find("meta", attrs={"name": "citation_journal_title"})
                if journal_tag:
                    metadata["journal"] = journal_tag.get("content", "")
                
                # Extract year
                date_tag = soup.find("meta", attrs={"name": "citation_publication_date"})
                if date_tag:
                    date_content = date_tag.get("content", "")
                    if date_content and len(date_content) >= 4:
                        metadata["year"] = date_content[:4]
                
                # Extract DOI
                doi_tag = soup.find("meta", attrs={"name": "citation_doi"})
                if doi_tag:
                    metadata["doi"] = doi_tag.get("content", "")
                
                return metadata
            except ImportError:
                # If BeautifulSoup is not available, use a simpler approach
                print("BeautifulSoup not available, using simple extraction")
                metadata = {
                    "title": "",
                    "authors": [],
                    "url": url
                }
                
                # Simple title extraction
                title_match = re.search(r"<title>(.*?)</title>", response.text)
                if title_match:
                    metadata["title"] = title_match.group(1)
                
                return metadata
        except Exception as e:
            print(f"Error extracting metadata from URL {url}: {e}")
            return None
    
    def _extract_from_nature_url(self, url: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a Nature URL"""
        # This would implement specific logic for Nature articles
        # For now, use the generic URL extractor
        return self._extract_from_generic_url(url)
    
    def _extract_from_science_url(self, url: str) -> Optional[Dict[str, Any]]:
        """Extract metadata from a Science URL"""
        # This would implement specific logic for Science articles
        # For now, use the generic URL extractor
        return self._extract_from_generic_url(url)
    
    def _parse_crossref_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Parse Crossref API response into standardized metadata format"""
        metadata = {
            "title": data.get("title", [""])[0] if isinstance(data.get("title", []), list) else data.get("title", ""),
            "authors": [],
            "journal": data.get("container-title", [""])[0] if isinstance(data.get("container-title", []), list) else "",
            "year": str(data.get("published", {}).get("date-parts", [[""]])[0][0]) if "published" in data else "",
            "volume": data.get("volume", ""),
            "issue": data.get("issue", ""),
            "pages": data.get("page", ""),
            "doi": data.get("DOI", ""),
            "url": data.get("URL", ""),
            "publisher": data.get("publisher", ""),
            "type": data.get("type", ""),
        }
        
        # Extract authors
        if "author" in data and isinstance(data["author"], list):
            for author in data["author"]:
                name = ""
                if "family" in author and "given" in author:
                    name = f"{author['family']}, {author['given']}"
                elif "name" in author:
                    name = author["name"]
                
                if name:
                    metadata["authors"].append(name)
        
        # Extract abstract if available
        if "abstract" in data:
            metadata["abstract"] = data["abstract"]
        
        # Extract keywords if available
        if "subject" in data and isinstance(data["subject"], list):
            metadata["keywords"] = data["subject"]
        
        return metadata
    
    def _parse_datacite_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Parse DataCite API response into standardized metadata format"""
        metadata = {
            "title": data.get("title", ""),
            "authors": [],
            "year": data.get("publicationYear", ""),
            "doi": data.get("doi", ""),
            "url": f"https://doi.org/{data.get('doi', '')}",
            "publisher": data.get("publisher", ""),
            "type": data.get("resourceType", {}).get("resourceTypeGeneral", ""),
        }
        
        # Extract authors
        if "creators" in data and isinstance(data["creators"], list):
            for creator in data["creators"]:
                if "name" in creator:
                    metadata["authors"].append(creator["name"])
        
        # Extract journal/container title if available
        if "container" in data and "title" in data["container"]:
            metadata["journal"] = data["container"]["title"]
        
        # Extract volume, issue, pages if available
        if "relatedIdentifiers" in data:
            for identifier in data["relatedIdentifiers"]:
                if identifier.get("relationType") == "IsPartOf":
                    parts = identifier.get("relatedIdentifier", "").split(":")
                    if len(parts) > 1:
                        if "volume" in parts[0].lower():
                            metadata["volume"] = parts[1]
                        elif "issue" in parts[0].lower():
                            metadata["issue"] = parts[1]
                        elif "page" in parts[0].lower():
                            metadata["pages"] = parts[1]
        
        return metadata
    
    def _parse_pubmed_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Parse PubMed API response into standardized metadata format"""
        metadata = {
            "title": data.get("title", ""),
            "authors": [],
            "journal": data.get("fulljournalname", "") or data.get("source", ""),
            "year": data.get("pubdate", "").split()[0] if data.get("pubdate", "") else "",
            "volume": data.get("volume", ""),
            "issue": data.get("issue", ""),
            "pages": data.get("pages", ""),
            "doi": next((id_data.get("value", "") for id_data in data.get("articleids", []) 
                         if id_data.get("idtype", "") == "doi"), ""),
            "pmid": next((id_data.get("value", "") for id_data in data.get("articleids", []) 
                         if id_data.get("idtype", "") == "pubmed"), ""),
            "url": f"https://pubmed.ncbi.nlm.nih.gov/{data.get('uid', '')}/" if data.get('uid', '') else "",
            "publisher": data.get("publishername", ""),
            "type": "journal-article",
        }
        
        # Extract authors
        if "authors" in data and isinstance(data["authors"], list):
            for author in data["authors"]:
                if "name" in author:
                    metadata["authors"].append(author["name"])
        
        # Extract abstract if available
        if "bookabstract" in data and data["bookabstract"]:
            metadata["abstract"] = data["bookabstract"]
        
        # Extract keywords if available
        if "keywords" in data and isinstance(data["keywords"], list):
            metadata["keywords"] = data["keywords"]
        
        return metadata
    
    def _parse_arxiv_entry(self, entry, ns: Dict[str, str]) -> Dict[str, Any]:
        """Parse arXiv API entry into standardized metadata format"""
        # Get the title
        title_element = entry.find("atom:title", ns)
        title = title_element.text if title_element is not None else ""
        
        # Get the authors
        authors = []
        for author_element in entry.findall("atom:author/atom:name", ns):
            # Convert author name to "Last, First" format
            name_parts = author_element.text.split()
            if len(name_parts) > 1:
                last_name = name_parts[-1]
                first_name = " ".join(name_parts[:-1])
                authors.append(f"{last_name}, {first_name}")
            else:
                authors.append(author_element.text)
        
        # Get the publication date
        date_element = entry.find("atom:published", ns)
        year = date_element.text[:4] if date_element is not None else ""
        
        # Get the DOI if available
        doi = ""
        for link_element in entry.findall("atom:link", ns):
            if link_element.get("title") == "doi":
                doi_url = link_element.get("href", "")
                if "doi.org/" in doi_url:
                    doi = doi_url.split("doi.org/")[-1]
        
        # Get the abstract
        summary_element = entry.find("atom:summary", ns)
        abstract = summary_element.text if summary_element is not None else ""
        
        # Get the URL
        url = ""
        for link_element in entry.findall("atom:link", ns):
            if link_element.get("rel") == "alternate":
                url = link_element.get("href", "")
        
        # Get the arXiv ID
        id_element = entry.find("atom:id", ns)
        arxiv_id = id_element.text.split("/")[-1] if id_element is not None else ""
        
        # Get the categories (subjects)
        categories = []
        for category_element in entry.findall("atom:category", ns):
            term = category_element.get("term")
            if term:
                categories.append(term)
        
        metadata = {
            "title": title,
            "authors": authors,
            "journal": "arXiv",
            "year": year,
            "doi": doi,
            "url": url,
            "arxiv_id": arxiv_id,
            "abstract": abstract,
            "type": "preprint",
            "keywords": categories
        }
        
        return metadata
    
    def _parse_openlibrary_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Parse Open Library API response into standardized metadata format"""
        metadata = {
            "title": data.get("title", ""),
            "authors": [],
            "year": "",
            "publisher": "",
            "isbn": "",
            "url": data.get("url", ""),
            "type": "book",
        }
        
        # Extract authors
        if "authors" in data and isinstance(data["authors"], list):
            for author in data["authors"]:
                if "name" in author:
                    metadata["authors"].append(author["name"])
        
        # Extract publication year
        if "publish_date" in data:
            publish_date = data["publish_date"]
            # Extract year from the publish date
            year_match = re.search(r"\d{4}", publish_date)
            if year_match:
                metadata["year"] = year_match.group(0)
        
        # Extract publisher
        if "publishers" in data and isinstance(data["publishers"], list) and len(data["publishers"]) > 0:
            metadata["publisher"] = data["publishers"][0].get("name", "")
        
        # Extract ISBN
        if "identifiers" in data and "isbn_13" in data["identifiers"]:
            metadata["isbn"] = data["identifiers"]["isbn_13"][0]
        elif "identifiers" in data and "isbn_10" in data["identifiers"]:
            metadata["isbn"] = data["identifiers"]["isbn_10"][0]
        
        return metadata


# ===== Parser Classes =====

class DOIParser:
    """Parse DOI to extract metadata"""
    
    def __init__(self, extractor):
        self.extractor = extractor
    
    def parse(self, doi: str) -> Optional[Dict[str, Any]]:
        """Parse a DOI and return metadata"""
        return self.extractor.extract_from_doi(doi)

class PMIDParser:
    """Parse PMID to extract metadata"""
    
    def __init__(self, extractor):
        self.extractor = extractor
    
    def parse(self, pmid: str) -> Optional[Dict[str, Any]]:
        """Parse a PMID and return metadata"""
        return self.extractor.extract_from_pmid(pmid)

class ArXivParser:
    """Parse arXiv IDs to extract metadata"""
    
    def __init__(self, extractor):
        self.extractor = extractor
    
    def parse(self, arxiv_id: str) -> Optional[Dict[str, Any]]:
        """Parse an arXiv ID and return metadata"""
        return self.extractor.extract_from_arxiv(arxiv_id)

class ISBNParser:
    """Parse ISBN numbers to extract metadata"""
    
    def __init__(self, extractor):
        self.extractor = extractor
    
    def parse(self, isbn: str) -> Optional[Dict[str, Any]]:
        """Parse an ISBN and return metadata"""
        return self.extractor.extract_from_isbn(isbn)

class URLParser:
    """Parse URLs to extract metadata"""
    
    def __init__(self, extractor):
        self.extractor = extractor
    
    def parse(self, url: str) -> Optional[Dict[str, Any]]:
        """Parse a URL and return metadata"""
        return self.extractor.extract_from_url(url)

class BibTexParser:
    """Parse BibTeX entries to extract metadata"""
    
    def parse(self, bibtex: str) -> Optional[Dict[str, Any]]:
        """Parse a BibTeX entry and return metadata"""
        try:
            # Simple BibTeX parser
            metadata = {
                "authors": []
            }
            
            # Extract entry type and key
            type_match = re.search(r'@(\w+)\s*\{([^,]+)', bibtex)
            if type_match:
                entry_type = type_match.group(1).lower()
                entry_key = type_match.group(2)
                
                metadata["type"] = entry_type
                metadata["id"] = entry_key
            
            # Extract fields
            field_pattern = re.compile(r'(\w+)\s*=\s*[{"]([^}"]*)["}\s]*[,]?')
            for field, value in field_pattern.findall(bibtex):
                field = field.lower()
                
                if field == "author":
                    # Split authors by "and"
                    authors = value.split(" and ")
                    for author in authors:
                        author = author.strip()
                        # Convert to "Last, First" format if needed
                        if "," not in author:
                            names = author.split()
                            if len(names) > 1:
                                first_names = " ".join(names[:-1])
                                last_name = names[-1]
                                author = f"{last_name}, {first_names}"
                        
                        metadata["authors"].append(author)
                else:
                    metadata[field] = value
            
            # Map common BibTeX fields to our standard format
            field_mapping = {
                "title": "title",
                "journal": "journal",
                "year": "year",
                "volume": "volume",
                "number": "issue",
                "pages": "pages",
                "doi": "doi",
                "url": "url",
                "publisher": "publisher",
                "booktitle": "booktitle",
                "address": "address",
                "abstract": "abstract",
                "keywords": "keywords"
            }
            
            standardized_metadata = {}
            for bibtex_field, standard_field in field_mapping.items():
                if bibtex_field in metadata:
                    standardized_metadata[standard_field] = metadata[bibtex_field]
            
            # Add authors
            standardized_metadata["authors"] = metadata["authors"]
            
            # Ensure essential fields are present
            for field in ["title", "authors", "year"]:
                if field not in standardized_metadata:
                    standardized_metadata[field] = ""
            
            return standardized_metadata
        except Exception as e:
            print(f"Error parsing BibTeX entry: {e}")
            return None

class RISParser:
    """Parse RIS (Research Information Systems) entries to extract metadata"""
    
    def parse(self, ris: str) -> Optional[Dict[str, Any]]:
        """Parse a RIS entry and return metadata"""
        try:
            metadata = {
                "authors": []
            }
            
            # Split the RIS string into lines
            lines = ris.strip().split("\n")
            
            # Process each line
            for line in lines:
                line = line.strip()
                if not line or len(line) < 6:
                    continue
                
                # Extract tag and value
                tag = line[:2].strip()
                value = line[6:].strip()
                
                if tag == "TY":
                    # Publication type
                    metadata["type"] = value
                elif tag == "AU" or tag == "A1" or tag == "A2":
                    # Author
                    author = value
                    # Convert to "Last, First" format if needed
                    if "," not in author:
                        names = author.split()
                        if len(names) > 1:
                            first_names = " ".join(names[:-1])
                            last_name = names[-1]
                            author = f"{last_name}, {first_names}"
                    
                    metadata["authors"].append(author)
                elif tag == "TI" or tag == "T1":
                    # Title
                    metadata["title"] = value
                elif tag == "JO" or tag == "JF" or tag == "JA":
                    # Journal
                    metadata["journal"] = value
                elif tag == "VL":
                    # Volume
                    metadata["volume"] = value
                elif tag == "IS":
                    # Issue
                    metadata["issue"] = value
                elif tag == "SP":
                    # Start page
                    if "pages" not in metadata:
                        metadata["pages"] = value
                    else:
                        metadata["pages"] = f"{value}-{metadata['pages']}"
                elif tag == "EP":
                    # End page
                    if "pages" not in metadata:
                        metadata["pages"] = f"-{value}"
                    else:
                        metadata["pages"] = f"{metadata['pages']}-{value}"
                elif tag == "PY" or tag == "Y1":
                    # Publication year
                    if "/" in value:
                        # Extract just the year
                        metadata["year"] = value.split("/")[0]
                    else:
                        metadata["year"] = value
                elif tag == "DO":
                    # DOI
                    metadata["doi"] = value
                elif tag == "UR":
                    # URL
                    metadata["url"] = value
                elif tag == "PB":
                    # Publisher
                    metadata["publisher"] = value
                elif tag == "AB":
                    # Abstract
                    metadata["abstract"] = value
                elif tag == "KW":
                    # Keywords
                    if "keywords" not in metadata:
                        metadata["keywords"] = []
                    metadata["keywords"].append(value)
            
            # Ensure essential fields are present
            for field in ["title", "authors", "year"]:
                if field not in metadata:
                    metadata[field] = ""
            
            return metadata
        except Exception as e:
            print(f"Error parsing RIS entry: {e}")
            return None


class IDDetector:
    """Detect and categorize different types of identifiers"""
    
    def detect_id_type(self, identifier: str) -> str:
        """
        Detect the type of identifier
        
        Args:
            identifier: The identifier string
            
        Returns:
            The detected type ("doi", "pmid", "arxiv", "isbn", "url", "bibtex", "ris", "unknown")
        """
        identifier = identifier.strip()
        
        # Check if it's a DOI
        if identifier.startswith("10.") or "doi.org" in identifier.lower():
            return "doi"
        
        # Check if it's a PMID
        if identifier.isdigit() and len(identifier) <= 8:
            return "pmid"
        
        # Check if it's an arXiv ID
        if (identifier.startswith("arXiv:") or 
            (re.match(r"\d{4}\.\d{4,5}", identifier)) or 
            (re.match(r"\d{7}", identifier) and len(identifier) == 7)):
            return "arxiv"
        
        # Check if it's an ISBN
        if (re.match(r"^(?:ISBN(?:-1[03])?:? )?(?=[0-9X]{10}$|(?=(?:[0-9]+[- ]){3})[- 0-9X]{13}$|97[89][0-9]{10}$|(?=(?:[0-9]+[- ]){4})[- 0-9]{17}$)(?:97[89][- ]?)?[0-9]{1,5}[- ]?[0-9]+[- ]?[0-9]+[- ]?[0-9X]$", identifier)):
            return "isbn"
        
        # Check if it's a URL
        if (identifier.startswith("http://") or 
            identifier.startswith("https://") or 
            identifier.startswith("www.")):
            return "url"
        
        # Check if it's a BibTeX entry
        if identifier.startswith("@") and "{" in identifier and "=" in identifier:
            return "bibtex"
        
        # Check if it's a RIS entry
        if "TY  -" in identifier and "ER  -" in identifier:
            return "ris"
        
        # If nothing matches, return unknown
        return "unknown"


# ===== Word Document Processing Classes =====

class CommandLineWordProcessor:
    """Process Word documents for citation management in command-line mode"""
    
    def __init__(self, citation_manager):
        """
        Initialize the word processor
        
        Args:
            citation_manager: The citation manager instance
        """
        self.citation_manager = citation_manager
        self.citation_pattern = r"\(\s*([^)]+)\s*,\s*(\d{4})\s*\)"  # Basic pattern for (Author, Year)
        
        # Additional patterns for different citation styles
        self.citation_patterns = {
            "author_year": r"\(\s*([^)]+)\s*,\s*(\d{4})\s*\)",  # (Author, Year)
            "superscript": r"(\w+)\s*(\d+)",  # Word1 for superscript citations
            "numbered": r"\[(\d+)\]",  # [1] for numbered citations
            "author_only": r"([A-Z][a-z]+)\s+et\s+al\.",  # Smith et al. for author only citations
        }
    
    def process_document(self, file_path: str, format_name: str, id_detector) -> Tuple[Document, Dict[str, Dict[str, Any]]]:
        """
        Process a Word document to extract citations
        
        Args:
            file_path: Path to the Word document
            format_name: Citation format name
            id_detector: ID detector instance
            
        Returns:
            Tuple of Document object and dictionary of citation metadata
        """
        print(f"Processing document: {file_path}")
        document = Document(file_path)
        citations = {}
        
        # Extract citations from paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text
            self._process_text_for_citations(text, citations, id_detector)
        
        # Extract citations from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_text_for_citations(paragraph.text, citations, id_detector)
        
        print(f"Found {len(citations)} potential citations")
        return document, citations
    
    def _process_text_for_citations(self, text: str, citations: Dict[str, Dict[str, Any]], id_detector) -> None:
        """
        Process text to extract citations
        
        Args:
            text: Text to process
            citations: Dictionary to store citations
            id_detector: ID detector instance
        """
        # Try all citation patterns
        for pattern_name, pattern in self.citation_patterns.items():
            matches = re.findall(pattern, text)
            for match in matches:
                if pattern_name == "author_year":
                    author, year = match
                    citation_key = f"{author.strip()}-{year.strip()}"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "author": author.strip(),
                            "year": year.strip(),
                            "pattern": pattern_name,
                            "source_text": f"({author.strip()}, {year.strip()})"
                        }
                elif pattern_name == "numbered":
                    number = match
                    citation_key = f"ref-{number}"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "number": number,
                            "pattern": pattern_name,
                            "source_text": f"[{number}]"
                        }
                elif pattern_name == "author_only":
                    author = match
                    citation_key = f"{author.strip()}-et-al"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "author": author.strip(),
                            "pattern": pattern_name,
                            "source_text": f"{author.strip()} et al."
                        }
        
        # Look for DOIs, PMIDs, arXiv IDs, and URLs in the text
        words = text.split()
        for word in words:
            word = word.strip(".,;()[]{}\"'")
            if word:
                id_type = id_detector.detect_id_type(word)
                if id_type != "unknown":
                    citation_key = f"{id_type}-{word}"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "id_type": id_type,
                            "id_value": word,
                            "pattern": "direct_id",
                            "source_text": word
                        }
    
    def update_document_with_citations(self, document: Document, 
                                  citations: Dict[str, Dict[str, Any]],
                                  extracted_metadata: Dict[str, Dict[str, Any]],
                                  format_name: str,
                                  output_path: str) -> bool:
        """
        Update document with formatted citations and add bibliography
        
        Args:
            document: The Word document object
            citations: Dictionary of extracted citation references
            extracted_metadata: Dictionary of extracted metadata
            format_name: Citation format name
            output_path: Path to save the updated document
            
        Returns:
            True if successful, False otherwise
        """
        # Create a mapping from citation keys to formatted citations
        formatted_citations = {}
        try:
            print(f"Updating document with {len(extracted_metadata)} citations in {format_name} format")
            
            # Get the formatter
            formatter = self.citation_manager.formatters.get(format_name.lower())
            if not formatter:
                print(f"Error: Formatter '{format_name}' not found")
                return False
            
            # Format each citation
            for citation_key, metadata in extracted_metadata.items():
                try:
                    formatted_in_text = formatter.format_in_text_citation(metadata)
                    formatted_citations[citation_key] = {
                        "in_text": formatted_in_text,
                        "bibliography": formatter.format_citation(metadata)
                    }
                    print(f"Formatted citation: {formatted_in_text} -> {formatter.format_citation(metadata)[:80]}...")
                except Exception as e:
                    print(f"Warning: Could not format citation {citation_key}: {e}")
                    continue
            
            # Update in-text citations - enhanced implementation
            try:
                print("Updating in-text citations...")
                citation_count = 0
                for paragraph in document.paragraphs:
                    changed = self._update_paragraph_citations(paragraph, citations, formatted_citations)
                    if changed:
                        citation_count += 1
                
                # Update citations in tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                changed = self._update_paragraph_citations(paragraph, citations, formatted_citations)
                                if changed:
                                    citation_count += 1
                                    
                print(f"Updated {citation_count} in-text citations")
            except Exception as e:
                print(f"Warning: Error updating in-text citations: {e}")
                # Continue processing even if in-text citations fail
        except Exception as e:
            print(f"Warning: Error in citation formatting: {e}")
            # Continue processing even if formatting fails
        
        # Create list of citations in proper order
        ordered_citations = []
        try:
            # Get unique metadata keys from citations
            used_metadata_keys = set()
            for citation in citations.values():
                if "metadata_key" in citation and citation["metadata_key"] in extracted_metadata:
                    used_metadata_keys.add(citation["metadata_key"])
            
            # Add any metadata keys that weren't explicitly linked to citations
            for key in extracted_metadata.keys():
                used_metadata_keys.add(key)
            
            # Create ordered list of bibliography entries
            ordered_citations = []
            for key in used_metadata_keys:
                if key in formatted_citations:
                    ordered_citations.append((key, formatted_citations[key]["bibliography"]))
            
            # Sort by first author's last name and year
            def get_sort_key(citation_tuple):
                key, bibliography = citation_tuple
                metadata = extracted_metadata.get(key, {})
                authors = metadata.get("authors", [])
                year = metadata.get("year", "9999")  # Default to high year if not found
                
                if authors:
                    # Get first author's last name
                    author = authors[0]
                    if "," in author:
                        last_name = author.split(",")[0].lower()
                    else:
                        parts = author.split()
                        last_name = parts[-1].lower() if parts else ""
                    return (last_name, year)
                return ("zzzzzz", year)  # Sort unknown authors at end
            
            # Sort the citations
            ordered_citations.sort(key=get_sort_key)
            
            print(f"Prepared {len(ordered_citations)} bibliography entries in alphabetical order")
        except Exception as e:
            print(f"Warning: Error organizing bibliography entries: {e}")
            # Fall back to using all available formatted citations
            ordered_citations = [(k, v["bibliography"]) for k, v in formatted_citations.items()]
            
        try:
            # Add bibliography section
            try:
                document.add_page_break()
            except Exception as e:
                print(f"Warning: Could not add page break: {e}")
                # Continue without page break
                
            # Create References title - with multiple fallback methods
            try:
                # Method 1: Try to add heading with style
                document.add_heading("References", level=1)
            except Exception as e:
                try:
                    # Method 2: Try to add a plain paragraph with bold text
                    print(f"Using simple heading instead of styled heading: {e}")
                    paragraph = document.add_paragraph("References")
                    paragraph.style = 'Normal'  # This will always exist
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(14)  # Approximation of heading size
                                    
                    # Try to set font size if possible
                    try:
                        from docx.shared import Pt
                        run.font.size = Pt(16)  # Approximate size for Heading 1
                    except ImportError:
                        # If docx.shared is not available, just use bold
                        pass
                except Exception as e2:
                    # Method 3: Last resort - just add plain text
                    print(f"Warning: Could not add bold heading: {e2}")
                    document.add_paragraph("REFERENCES")
                    
            # Add each citation to the bibliography
            if ordered_citations:
                print(f"Adding {len(ordered_citations)} citations to bibliography")
                for i, (citation_key, bibliography) in enumerate(ordered_citations, 1):
                    try:
                        paragraph = document.add_paragraph()
                        # Add number for numbered citations
                        paragraph.text = f"{i}. {bibliography}"
                        print(f"Added bibliography entry {i}: {bibliography[:80]}...")
                    except Exception as e:
                        print(f"Warning: Could not add bibliography entry for {citation_key}: {e}")
                        continue
            else:
                print("No citations to add to bibliography")
        except Exception as e:
            print(f"Warning: Error adding bibliography: {e}")
            # Continue without bibliography if it fails
        
        # Save the document - this is the most critical part
        try:
            document.save(output_path)
            print(f"Document saved to {output_path}")
            return True
        except Exception as e:
            print(f"Error saving document to {output_path}: {e}")
            
            # Try an alternative approach to saving if the first one fails
            try:
                # Create a new document and copy content
                from docx import Document as NewDocument
                new_doc = NewDocument()
                
                # Copy the paragraphs from the original document
                for para in document.paragraphs:
                    new_para = new_doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                
                # Save the new document
                new_doc.save(output_path)
                print(f"Document saved using alternative method to {output_path}")
                return True
            except Exception as e2:
                print(f"Fatal error: Could not save document using alternative method: {e2}")
                return False
    
    def _update_paragraph_citations(self, paragraph, citations, formatted_citations) -> bool:
        """
        Update citations in a paragraph
        
        Args:
            paragraph: The paragraph object
            citations: Dictionary of citation references
            formatted_citations: Dictionary of formatted citations
        
        Returns:
            True if paragraph was modified, False otherwise
        """
        text = paragraph.text
        updated_text = text
        modified = False
        
        # Replace each citation with the formatted version
        for citation_key, citation in citations.items():
            if "source_text" in citation and citation["source_text"] in text:
                # Find the metadata key for this citation
                if "metadata_key" in citation and citation["metadata_key"] in formatted_citations:
                    metadata_key = citation["metadata_key"]
                    formatted_citation = formatted_citations[metadata_key]["in_text"]
                    
                    # Only replace if we have something valid
                    if formatted_citation:
                        updated_text = updated_text.replace(
                            citation["source_text"], 
                            formatted_citation
                        )
                        modified = True
                else:
                    # Try to find a direct match in formatted_citations
                    for fmt_key, fmt_citation in formatted_citations.items():
                        if citation_key in fmt_key:
                            updated_text = updated_text.replace(
                                citation["source_text"],
                                fmt_citation["in_text"]
                            )
                            modified = True
                            break
        
        # Update the paragraph text if changes were made
        if updated_text != text:
            paragraph.text = updated_text
            return True
        
        return False
    
    def _get_matching_key(self, citation, formatted_citations) -> Optional[str]:
        """
        Find a matching key in the formatted citations
        
        Args:
            citation: Citation reference
            formatted_citations: Dictionary of formatted citations
            
        Returns:
            Matching key if found, None otherwise
        """
        # This is a simplified matching logic
        # A more sophisticated implementation would use fuzzy matching
        
        # Direct match by ID
        if "id_type" in citation and "id_value" in citation:
            for key in formatted_citations:
                if citation["id_type"] in key and citation["id_value"] in key:
                    return key
        
        # Match by author and year
        if "author" in citation and "year" in citation:
            author = citation["author"].lower()
            year = citation["year"]
            
            for key in formatted_citations:
                metadata_key = key.split("-metadata-")[0] if "-metadata-" in key else key
                if author in metadata_key.lower() and year in metadata_key:
                    return key
        
        return None

class NumberedCitationProcessor:
    """Process citations with a numbered system like [1], [2], etc."""
    
    def __init__(self, citation_manager):
        """Initialize with citation manager"""
        self.citation_manager = citation_manager
        self.citations_order = {}  # Maps metadata keys to citation numbers
        self.current_number = 1
    
    def process_document(self, document, citations, extracted_metadata, format_name):
        """
        First pass to collect and number all citations
        
        Args:
            document: Word document
            citations: Dictionary of extracted citation references
            extracted_metadata: Metadata dictionary
            format_name: Citation format name
            
        Returns:
            Dictionary mapping citation keys to citation numbers
        """
        # Reset numbering
        self.citations_order = {}
        self.current_number = 1
        
        # First, get all unique citations
        for paragraph in document.paragraphs:
            self._collect_citations_from_text(paragraph.text, citations)
            
        # Then for tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._collect_citations_from_text(paragraph.text, citations)
        
        print(f"Found {len(self.citations_order)} unique citations for numbering")
        return self.citations_order
    
    def _collect_citations_from_text(self, text, citations):
        """Collect citations from text and assign numbers"""
        for citation_key, citation in citations.items():
            if "source_text" in citation and citation["source_text"] in text:
                # Check if we have metadata for this citation
                if "metadata_key" in citation and citation["metadata_key"] not in self.citations_order:
                    self.citations_order[citation["metadata_key"]] = self.current_number
                    self.current_number += 1
    
    def format_in_text_citations(self, document, citations, citation_numbers):
        """
        Replace in-text citations with numbered format [1], [2], etc.
        
        Args:
            document: Word document
            citations: Dictionary of citation references
            citation_numbers: Dictionary mapping metadata keys to numbers
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Process paragraphs
            for paragraph in document.paragraphs:
                self._update_paragraph_with_numbers(paragraph, citations, citation_numbers)
            
            # Process tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._update_paragraph_with_numbers(paragraph, citations, citation_numbers)
            
            return True
        except Exception as e:
            print(f"Error formatting numbered citations: {e}")
            return False
    
    def _update_paragraph_with_numbers(self, paragraph, citations, citation_numbers):
        """Update paragraph with numbered citations"""
        text = paragraph.text
        updated_text = text
        
        # Replace each citation with its number
        for citation_key, citation in citations.items():
            if "source_text" in citation and citation["source_text"] in text:
                if "metadata_key" in citation and citation["metadata_key"] in citation_numbers:
                    number = citation_numbers[citation["metadata_key"]]
                    updated_text = updated_text.replace(
                        citation["source_text"],
                        f"[{number}]"
                    )
        
        # Update the paragraph text if changes were made
        if updated_text != text:
            paragraph.text = updated_text
    
    def generate_numbered_bibliography(self, document, extracted_metadata, citation_numbers, format_name):
        """
        Generate a numbered bibliography in the document
        
        Args:
            document: Word document
            extracted_metadata: Dictionary of extracted metadata
            citation_numbers: Dictionary mapping metadata keys to numbers
            format_name: Citation format name
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Add bibliography section
            document.add_page_break()
            document.add_heading("References", level=1)
            
            # Get formatter
            formatter = self.citation_manager.formatters.get(format_name.lower())
            if not formatter:
                print(f"Error: Formatter '{format_name}' not found")
                return False
            
            # Create ordered list of citations
            ordered_citations = []
            for metadata_key, number in citation_numbers.items():
                if metadata_key in extracted_metadata:
                    metadata = extracted_metadata[metadata_key]
                    bibliography = formatter.format_citation(metadata)
                    ordered_citations.append((number, bibliography))
            
            # Sort by citation number
            ordered_citations.sort()
            
            # Add each citation to the bibliography
            for number, bibliography in ordered_citations:
                paragraph = document.add_paragraph()
                paragraph.text = f"{number}. {bibliography}"
            
            return True
        except Exception as e:
            print(f"Error generating numbered bibliography: {e}")
            return False
class EnhancedCitationManager:
    """Enhanced command-line citation manager with support for multiple sources and formats"""
    
    def __init__(self):
        """Initialize the citation manager"""
        self.metadata_extractor = EnhancedMetadataExtractor()
        
        # Initialize parsers
        self.parsers = {
            "doi": DOIParser(self.metadata_extractor),
            "pmid": PMIDParser(self.metadata_extractor),
            "arxiv": ArXivParser(self.metadata_extractor),
            "isbn": ISBNParser(self.metadata_extractor),
            "url": URLParser(self.metadata_extractor),
            "bibtex": BibTexParser(),
            "ris": RISParser()
        }
        
        # Initialize formatters
        self.formatters = {
            # Basic styles
            "elsevier": ElsevierFormatter(),
            "springer": SpringerFormatter(),
            "apa": APAFormatter(),
            
            # General journal styles
            "nature": NatureFormatter(),
            "science": ScienceFormatter(),
            "ieee": IEEEFormatter(),
            "chicago": ChicagoFormatter(),
            "mla": MLAFormatter(),
            "harvard": HarvardFormatter(),
            "vancouver": VancouverFormatter(),
            "bmc": BMCFormatter(),
            "plos": PLOSFormatter(),
            "cell": CellFormatter(),
            
            # Medical journal styles
            "jama": JAMAFormatter(),
            "bmj": BMJFormatter(),
            "nejm": NEJMFormatter(),
            
            # Scientific journal styles
            "jbc": JBCFormatter(),
            "rsc": RSCFormatter(),
            "acs": ACSFormatter(),
            "aip": AIPFormatter(),
            
            # Computer science and humanities
            "acm": ACMFormatter(),
            "oxford": OxfordFormatter()
        }
        
        self.citation_data = {}
        self.id_detector = IDDetector()
    
    def extract_metadata(self, id_type: str, id_value: str) -> Optional[Dict[str, Any]]:
        """
        Extract metadata based on ID type and value
        
        Args:
            id_type: Type of identifier
            id_value: Identifier value
            
        Returns:
            Metadata dictionary if successful, None otherwise
        """
        parser = self.parsers.get(id_type.lower())
        if parser:
            return parser.parse(id_value)
        else:
            print(f"Unsupported ID type: {id_type}")
            return None
    
    def format_citation(self, metadata: Dict[str, Any], format_name: str) -> Optional[str]:
        """
        Format metadata using the specified formatter
        
        Args:
            metadata: Metadata dictionary
            format_name: Citation format name
            
        Returns:
            Formatted citation if successful, None otherwise
        """
        formatter = self.formatters.get(format_name.lower())
        if formatter:
            return formatter.format_citation(metadata)
        else:
            print(f"Unsupported format: {format_name}")
            return None
    
    def process_single_citation(self, id_type: str, id_value: str, format_name: str) -> Optional[str]:
        """
        Process a single citation
        
        Args:
            id_type: Type of identifier
            id_value: Identifier value
            format_name: Citation format name
            
        Returns:
            Formatted citation if successful, None otherwise
        """
        # Auto-detect ID type if set to "auto"
        if id_type.lower() == "auto":
            id_type = self.id_detector.detect_id_type(id_value)
            print(f"Detected ID type: {id_type}")
        
        metadata = self.extract_metadata(id_type, id_value)
        if metadata:
            citation = self.format_citation(metadata, format_name)
            if citation:
                # Store metadata for later use
                key = f"{id_type}_{id_value}"
                metadata["formatted_citation"] = citation
                self.citation_data[key] = metadata
                return citation
        return None
    
    def process_batch_citations(self, ids: List[str], id_type: str, format_name: str) -> List[str]:
        """
        Process multiple citations
        
        Args:
            ids: List of identifiers
            id_type: Type of identifier
            format_name: Citation format name
            
        Returns:
            List of formatted citations
        """
        results = []
        for id_value in ids:
            id_value = id_value.strip()
            if not id_value:
                continue
            
            # Auto-detect ID type if set to "auto"
            current_id_type = id_type
            if id_type.lower() == "auto":
                current_id_type = self.id_detector.detect_id_type(id_value)
                print(f"Detected ID type for '{id_value}': {current_id_type}")
            
            citation = self.process_single_citation(current_id_type, id_value, format_name)
            if citation:
                results.append(citation)
            else:
                results.append(f"Failed to retrieve/format citation for {current_id_type} {id_value}")
        
        return results
    
    def process_file(self, file_path: str, id_type: str, format_name: str) -> List[str]:
        """
        Process IDs from a file
        
        Args:
            file_path: Path to the file containing identifiers
            id_type: Type of identifier
            format_name: Citation format name
            
        Returns:
            List of formatted citations
        """
        try:
            with open(file_path, "r") as file:
                ids = [line.strip() for line in file if line.strip()]
            return self.process_batch_citations(ids, id_type, format_name)
        except Exception as e:
            print(f"Error processing file: {e}")
            return []
    
    def export_citations(self, citations: List[str], output_path: str) -> bool:
        """
        Export citations to a file
        
        Args:
            citations: List of formatted citations
            output_path: Path to save the citations
            
        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_path, "w") as file:
                for citation in citations:
                    file.write(citation + "\n\n")
            return True
        except Exception as e:
            print(f"Error exporting citations: {e}")
            return False
    
    def export_bibtex(self, output_path: str) -> bool:
        """
        Export citations to a BibTeX file
        
        Args:
            output_path: Path to save the BibTeX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_path, "w") as file:
                for key, metadata in self.citation_data.items():
                    bibtex_entry = self._convert_to_bibtex(key, metadata)
                    file.write(bibtex_entry + "\n\n")
            return True
        except Exception as e:
            print(f"Error exporting BibTeX: {e}")
            return False
    
    def _convert_to_bibtex(self, key: str, metadata: Dict[str, Any]) -> str:
        """
        Convert metadata to BibTeX format
        
        Args:
            key: Citation key
            metadata: Metadata dictionary
            
        Returns:
            BibTeX entry
        """
        # Determine entry type
        entry_type = "article"
        if "type" in metadata:
            if metadata["type"] == "book":
                entry_type = "book"
            elif metadata["type"] == "inproceedings" or metadata["type"] == "conference-paper":
                entry_type = "inproceedings"
            elif metadata["type"] == "preprint":
                entry_type = "unpublished"
        
        # Create BibTeX key
        if "authors" in metadata and metadata["authors"]:
            first_author = metadata["authors"][0].split(",")[0]
            bibtex_key = f"{first_author}{metadata.get('year', '')}"
        else:
            bibtex_key = f"citation{key}"
        
        # Start building the BibTeX entry
        bibtex = f"@{entry_type}{{{bibtex_key},\n"
        
        # Add authors
        if "authors" in metadata and metadata["authors"]:
            authors = " and ".join(metadata["authors"])
            bibtex += f"  author = {{{authors}}},\n"
        
        # Add title
        if "title" in metadata and metadata["title"]:
            bibtex += f"  title = {{{metadata['title']}}},\n"
        
        # Add journal/booktitle
        if "journal" in metadata and metadata["journal"]:
            if entry_type == "inproceedings":
                bibtex += f"  booktitle = {{{metadata['journal']}}},\n"
            else:
                bibtex += f"  journal = {{{metadata['journal']}}},\n"
        
        # Add year
        if "year" in metadata and metadata["year"]:
            bibtex += f"  year = {{{metadata['year']}}},\n"
        
        # Add volume
        if "volume" in metadata and metadata["volume"]:
            bibtex += f"  volume = {{{metadata['volume']}}},\n"
        
        # Add number/issue
        if "issue" in metadata and metadata["issue"]:
            bibtex += f"  number = {{{metadata['issue']}}},\n"
        
        # Add pages
        if "pages" in metadata and metadata["pages"]:
            bibtex += f"  pages = {{{metadata['pages']}}},\n"
        
        # Add DOI
        if "doi" in metadata and metadata["doi"]:
            bibtex += f"  doi = {{{metadata['doi']}}},\n"
        
        # Add URL
        if "url" in metadata and metadata["url"]:
            bibtex += f"  url = {{{metadata['url']}}},\n"
        
        # Add publisher
        if "publisher" in metadata and metadata["publisher"]:
            bibtex += f"  publisher = {{{metadata['publisher']}}},\n"
        
        # Add abstract
        if "abstract" in metadata and metadata["abstract"]:
            # Truncate long abstracts
            abstract = metadata["abstract"]
            if len(abstract) > 500:
                abstract = abstract[:497] + "..."
            bibtex += f"  abstract = {{{abstract}}},\n"
        
        # Remove the trailing comma and newline, and close the entry
        bibtex = bibtex.rstrip(",\n") + "\n}"
        
        return bibtex

# ===== Command-Line Interface =====

def main():
    """Main entry point for the citation manager"""
    parser = argparse.ArgumentParser(description="Enhanced Citation Manager")
    
    # Create subparsers for different commands
    subparsers = parser.add_subparsers(dest="command", help="Command to execute")
    
    # Define available formats
    format_choices = [
        "elsevier", "springer", "apa", "nature", "science", "ieee", 
        "chicago", "mla", "harvard", "vancouver", "bmc", "plos", "cell",
        "jama", "bmj", "nejm", "jbc", "rsc", "acs", "aip", "acm", "oxford"
    ]
    
    # Define available ID types
    id_type_choices = ["doi", "pmid", "arxiv", "isbn", "url", "bibtex", "ris", "auto"]
    
    # Single citation parser
    single_parser = subparsers.add_parser("single", help="Generate a single citation")
    single_parser.add_argument("--id-type", "-t", choices=id_type_choices, default="auto", 
                              help="Type of identifier")
    single_parser.add_argument("--id", "-i", required=True, help="Identifier value")
    single_parser.add_argument("--format", "-f", choices=format_choices, 
                              default="apa", help="Citation format")
    single_parser.add_argument("--output", "-o", help="Output file path")
    single_parser.add_argument("--bibtex", "-b", help="Export as BibTeX to file")
    
    # Batch citation parser
    batch_parser = subparsers.add_parser("batch", help="Process multiple citations")
    batch_parser.add_argument("--id-type", "-t", choices=id_type_choices, default="auto", 
                             help="Type of identifier")
    batch_parser.add_argument("--ids", "-i", nargs="+", help="List of identifiers")
    batch_parser.add_argument("--file", "-l", help="File containing identifiers (one per line)")
    batch_parser.add_argument("--format", "-f", choices=format_choices, 
                             default="apa", help="Citation format")
    batch_parser.add_argument("--output", "-o", help="Output file path")
    batch_parser.add_argument("--bibtex", "-b", help="Export as BibTeX to file")
    
    # Word document parser
    word_parser = subparsers.add_parser("word", help="Process Word documents")
    word_parser.add_argument("--input", "-i", required=True, help="Input Word document")
    word_parser.add_argument("--output", "-o", required=True, help="Output Word document")
    word_parser.add_argument("--format", "-f", choices=format_choices, default="apa", 
                           help="Citation format")
    word_parser.add_argument("--id-list", "-l", help="File containing identifiers for the bibliography")
    word_parser.add_argument("--id-type", "-t", choices=id_type_choices, 
                           default="auto", help="Type of identifiers in the ID list")
    
    # Interactive mode parser
    interactive_parser = subparsers.add_parser("interactive", help="Run in interactive mode")
    
    # Parse arguments
    args = parser.parse_args()
    
    # Create citation manager
    citation_manager = EnhancedCitationManager()
    
    if args.command == "single":
        # Process a single citation
        citation = citation_manager.process_single_citation(
            args.id_type, args.id, args.format
        )
        
        if citation:
            print("\nGenerated Citation:")
            print(citation)
            
            if args.output:
                if citation_manager.export_citations([citation], args.output):
                    print(f"\nCitation exported to {args.output}")
                else:
                    print(f"\nFailed to export citation to {args.output}")
            
            if args.bibtex:
                if citation_manager.export_bibtex(args.bibtex):
                    print(f"\nBibTeX exported to {args.bibtex}")
                else:
                    print(f"\nFailed to export BibTeX to {args.bibtex}")
        else:
            print(f"\nFailed to generate citation for {args.id_type} {args.id}")
    
    elif args.command == "batch":
        # Process multiple citations
        if args.file:
            citations = citation_manager.process_file(args.file, args.id_type, args.format)
        elif args.ids:
            citations = citation_manager.process_batch_citations(args.ids, args.id_type, args.format)
        else:
            print("Error: Either --ids or --file must be provided for batch processing")
            return
        
        if citations:
            print(f"\nGenerated {len(citations)} citations:")
            for i, citation in enumerate(citations, 1):
                print(f"\n{i}. {citation}")
            
            if args.output:
                if citation_manager.export_citations(citations, args.output):
                    print(f"\nCitations exported to {args.output}")
                else:
                    print(f"\nFailed to export citations to {args.output}")
            
            if args.bibtex:
                if citation_manager.export_bibtex(args.bibtex):
                    print(f"\nBibTeX exported to {args.bibtex}")
                else:
                    print(f"\nFailed to export BibTeX to {args.bibtex}")
        else:
            print("\nNo citations generated")
    
    elif args.command == "word":
        # Process Word document
        handle_word_command(args, citation_manager)
    
    elif args.command == "interactive":
        # Run in interactive mode
        run_interactive_mode(citation_manager, format_choices, id_type_choices)
    
    else:
        # Show help if no command is provided
        parser.print_help()

# Add this to the handle_word_command function in complete_citation_manager.py

def handle_word_command(args, citation_manager):
    """Handle the word document processing command"""
    if not args.input or not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found")
        return
    
    # Check if python-docx is installed
    try:
        import docx
    except ImportError:
        print("Error: python-docx package is required for Word document processing")
        print("Please install it with: pip install python-docx")
        return
    
    print(f"Processing Word document: {args.input}")
    print(f"Citation format: {args.format}")
    
    # Create ID detector
    id_detector = IDDetector()
    
    # Create Word processor
    word_processor = CommandLineWordProcessor(citation_manager)
    
    # Process the document
    document, citations = word_processor.process_document(args.input, args.format, id_detector)
    
    # Extract metadata for citations
    print("Extracting metadata for citations...")
    extracted_metadata = {}
    citation_lookup = {}  # Map source_text to metadata keys
    
    # Load identifiers from file if provided
    identifiers = []
    if args.id_list and os.path.exists(args.id_list):
        with open(args.id_list, "r") as file:
            identifiers = [line.strip() for line in file if line.strip()]
        
        for identifier in identifiers:
            if args.id_type == "auto":
                id_type = id_detector.detect_id_type(identifier)
            else:
                id_type = args.id_type
            
            print(f"Processing identifier: {identifier} (type: {id_type})")
            metadata = citation_manager.extract_metadata(id_type, identifier)
            
            if metadata:
                key = f"{id_type}-{identifier}-metadata"
                extracted_metadata[key] = metadata
                
                # Create default source text for this citation
                if "authors" in metadata and metadata["authors"] and "year" in metadata:
                    first_author = metadata["authors"][0].split(",")[0]
                    year = metadata["year"]
                    if len(metadata["authors"]) > 1:
                        source_text = f"{first_author} et al., {year}"
                    else:
                        source_text = f"{first_author}, {year}"
                    citation_lookup[source_text.lower()] = key
                
                print(f"  Extracted metadata for {id_type} {identifier}")
    
    # Extract metadata from citations found in the document
    for citation_key, citation in citations.items():
        if "id_type" in citation and "id_value" in citation:
            id_type = citation["id_type"]
            id_value = citation["id_value"]
            
            print(f"Processing direct identifier in document: {id_type} {id_value}")
            metadata = citation_manager.extract_metadata(id_type, id_value)
            
            if metadata:
                key = f"{id_type}-{id_value}-metadata"
                extracted_metadata[key] = metadata
                citation_lookup[citation["source_text"].lower()] = key
                print(f"  Extracted metadata for {id_type} {id_value}")
    
    # Match author-year citations with metadata
    for citation_key, citation in citations.items():
        # Skip citations we've already processed
        if "source_text" in citation and citation["source_text"].lower() in citation_lookup:
            continue
            
        if "pattern" in citation and citation["pattern"] == "author_year":
            author = citation["author"].lower()
            year = citation["year"]
            source_text = citation["source_text"].lower()
            
            print(f"Processing author-year citation: {citation['source_text']}")
            
            # First check if we already have this exact source text
            if source_text in citation_lookup:
                continue
                
            # Try to find a matching metadata entry
            matched = False
            for metadata_key, metadata in extracted_metadata.items():
                if "authors" in metadata and metadata["authors"] and "year" in metadata:
                    # Get the first author's last name
                    first_author = metadata["authors"][0].split(",")[0].lower()
                    metadata_year = metadata["year"]
                    
                    # Check if this metadata matches our citation
                    if (first_author in author or author in first_author) and year == metadata_year:
                        citation_lookup[source_text] = metadata_key
                        print(f"  Matched citation {citation_key} with metadata {metadata_key}")
                        matched = True
                        break
            
            # If no match found, try to fetch metadata for this author/year
            if not matched:
                print(f"  No metadata match found for {citation['source_text']}, attempting to search...")
                try:
                    # Attempt to search for this publication
                    search_query = f"{citation['author']} {citation['year']}"
                    
                    # Try to make a naive DOI search via Crossref
                    search_url = f"https://api.crossref.org/works?query={search_query.replace(' ', '+')}&rows=1"
                    response = requests.get(search_url)
                    data = response.json()
                    
                    if "message" in data and "items" in data["message"] and data["message"]["items"]:
                        item = data["message"]["items"][0]
                        if "DOI" in item:
                            doi = item["DOI"]
                            print(f"  Found potential DOI: {doi}")
                            metadata = citation_manager.metadata_extractor.extract_from_doi(doi)
                            
                            if metadata:
                                key = f"doi-{doi}-metadata"
                                extracted_metadata[key] = metadata
                                citation_lookup[source_text] = key
                                print(f"  Successfully retrieved metadata for {citation['source_text']} via DOI search")
                            else:
                                print(f"  Failed to extract metadata from found DOI: {doi}")
                        else:
                            print("  Found search result but no DOI available")
                    else:
                        print("  No search results found")
                        
                    # Even if we didn't find metadata, create a placeholder
                    if source_text not in citation_lookup:
                        # Create minimal metadata from the citation itself
                        minimal_metadata = {
                            "title": f"[No title found for {citation['source_text']}]",
                            "authors": [f"{citation['author'].strip()}, "],
                            "year": citation["year"],
                            "journal": "[Journal not found]"
                        }
                        key = f"placeholder-{citation_key}"
                        extracted_metadata[key] = minimal_metadata
                        citation_lookup[source_text] = key
                        print(f"  Created placeholder metadata for {citation['source_text']}")
                        
                except Exception as e:
                    print(f"  Error searching for metadata: {e}")
                    # Create minimal metadata as a fallback
                    minimal_metadata = {
                        "title": f"[No title found for {citation['source_text']}]",
                        "authors": [f"{citation['author'].strip()}, "],
                        "year": citation["year"],
                        "journal": "[Journal not found]"
                    }
                    key = f"placeholder-{citation_key}"
                    extracted_metadata[key] = minimal_metadata
                    citation_lookup[source_text] = key
                    print(f"  Created placeholder metadata for {citation['source_text']}")
    
    # Link citations to their metadata using the lookup table
    for citation_key, citation in citations.items():
        if "source_text" in citation:
            source_text = citation["source_text"].lower()
            if source_text in citation_lookup:
                metadata_key = citation_lookup[source_text]
                citation["metadata_key"] = metadata_key
                print(f"Linked citation '{citation['source_text']}' to metadata key '{metadata_key}'")
    
    # Update the document with citations
    if extracted_metadata:
        # Print a summary of what will be updated
        print("\nFound the following references:")
        for i, (key, metadata) in enumerate(extracted_metadata.items(), 1):
            formatter = citation_manager.formatters.get(args.format.lower())
            if formatter:
                citation_str = formatter.format_citation(metadata)
                author_str = ""
                if metadata.get("authors"):
                    if len(metadata["authors"]) == 1:
                        author_str = metadata["authors"][0].split(",")[0]
                    else:
                        author_str = metadata["authors"][0].split(",")[0] + " et al."
                year_str = metadata.get("year", "")
                print(f"{i}. {author_str} ({year_str}) - Will be formatted as full citation")
            
    #     # Proceed with document update
    #     word_processor.update_document_with_citations(
    #         document, citations, extracted_metadata, args.format, args.output
    #     )
    # else:
    #     print("No metadata extracted. Cannot update the document.")
        if args.format.lower() in ["vancouver", "ieee"]:
            print("Using numbered citation style...")
            numbered_processor = NumberedCitationProcessor(citation_manager)
            citation_numbers = numbered_processor.process_document(document, citations, extracted_metadata, args.format)
            
            # Replace in-text citations with numbers
            numbered_processor.format_in_text_citations(document, citations, citation_numbers)
            
            # Generate numbered bibliography
            numbered_processor.generate_numbered_bibliography(document, extracted_metadata, citation_numbers, args.format)
            
            # Save document
            try:
                document.save(args.output)
                print(f"Document saved to {args.output} with numbered citations")
                return
            except Exception as e:
                print(f"Error saving document with numbered citations: {e}")
                # Continue to standard processing as fallback
        else:
            # Standard processing for non-numbered styles
            word_processor.update_document_with_citations(
                document, citations, extracted_metadata, args.format, args.output
            )
    else:
        print("No metadata extracted. Cannot update the document.")


def run_interactive_mode(citation_manager, format_choices, id_type_choices):
    """Run the citation manager in interactive mode"""
    print("\nEnhanced Citation Manager - Interactive Mode")
    print("===========================================")
    
    while True:
        print("\nOptions:")
        print("1. Generate a single citation")
        print("2. Process multiple citations")
        print("3. Process citations from a file")
        print("4. Process a Word document")
        print("5. Exit")
        
        choice = input("\nEnter your choice (1-5): ")
        
        if choice == "1":
            # Single citation
            id_type = get_valid_input(
                f"Enter ID type ({'/'.join(id_type_choices)}): ", 
                lambda x: x.lower() in id_type_choices,
                "auto"
            )
            
            id_value = input("Enter ID value: ")
            
            format_name = get_valid_input(
                f"Enter citation format ({'/'.join(format_choices)}): ", 
                lambda x: x.lower() in format_choices,
                "apa"
            )
            
            print("\nProcessing citation...")
            citation = citation_manager.process_single_citation(id_type, id_value, format_name)
            
            if citation:
                print("\nGenerated Citation:")
                print(citation)
                
                save_option = input("\nSave to file? (y/n): ")
                if save_option.lower() == "y":
                    output_path = input("Enter output file path: ")
                    citation_manager.export_citations([citation], output_path)
                    print(f"Citation saved to {output_path}")
                
                bibtex_option = input("\nExport as BibTeX? (y/n): ")
                if bibtex_option.lower() == "y":
                    bibtex_path = input("Enter BibTeX file path: ")
                    citation_manager.export_bibtex(bibtex_path)
                    print(f"BibTeX exported to {bibtex_path}")
            else:
                print(f"\nFailed to generate citation for {id_type} {id_value}")
        
        elif choice == "2":
            # Batch citations
            id_type = get_valid_input(
                f"Enter ID type ({'/'.join(id_type_choices)}): ", 
                lambda x: x.lower() in id_type_choices,
                "auto"
            )
            
            print("Enter IDs (one per line, empty line to finish):")
            ids = []
            while True:
                id_line = input()
                if not id_line:
                    break
                ids.append(id_line)
            
            if not ids:
                print("No IDs entered")
                continue
            
            format_name = get_valid_input(
                f"Enter citation format ({'/'.join(format_choices)}): ", 
                lambda x: x.lower() in format_choices,
                "apa"
            )
            
            print("\nProcessing citations...")
            citations = citation_manager.process_batch_citations(ids, id_type, format_name)
            
            if citations:
                print(f"\nGenerated {len(citations)} citations:")
                for i, citation in enumerate(citations, 1):
                    print(f"\n{i}. {citation}")
                
                save_option = input("\nSave to file? (y/n): ")
                if save_option.lower() == "y":
                    output_path = input("Enter output file path: ")
                    citation_manager.export_citations(citations, output_path)
                    print(f"Citations saved to {output_path}")
                
                bibtex_option = input("\nExport as BibTeX? (y/n): ")
                if bibtex_option.lower() == "y":
                    bibtex_path = input("Enter BibTeX file path: ")
                    citation_manager.export_bibtex(bibtex_path)
                    print(f"BibTeX exported to {bibtex_path}")
            else:
                print("\nNo citations generated")
        
        elif choice == "3":
            # Process from file
            id_type = get_valid_input(
                f"Enter ID type ({'/'.join(id_type_choices)}): ", 
                lambda x: x.lower() in id_type_choices,
                "auto"
            )
            
            file_path = input("Enter file path: ")
            if not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                continue
            
            format_name = get_valid_input(
                f"Enter citation format ({'/'.join(format_choices)}): ", 
                lambda x: x.lower() in format_choices,
                "apa"
            )
            
            print("\nProcessing citations from file...")
            citations = citation_manager.process_file(file_path, id_type, format_name)
            
            if citations:
                print(f"\nGenerated {len(citations)} citations:")
                for i, citation in enumerate(citations, 1):
                    print(f"\n{i}. {citation}")
                
                save_option = input("\nSave to file? (y/n): ")
                if save_option.lower() == "y":
                    output_path = input("Enter output file path: ")
                    citation_manager.export_citations(citations, output_path)
                    print(f"Citations saved to {output_path}")
                
                bibtex_option = input("\nExport as BibTeX? (y/n): ")
                if bibtex_option.lower() == "y":
                    bibtex_path = input("Enter BibTeX file path: ")
                    citation_manager.export_bibtex(bibtex_path)
                    print(f"BibTeX exported to {bibtex_path}")
            else:
                print("\nNo citations generated")
        
        elif choice == "4":
            # Process Word document
            try:
                from docx import Document
            except ImportError:
                print("Error: python-docx package is required for Word document processing")
                print("Please install it with: pip install python-docx")
                continue
            
            input_path = input("Enter input Word document path: ")
            if not os.path.exists(input_path):
                print(f"File not found: {input_path}")
                continue
            
            output_path = input("Enter output Word document path: ")
            
            format_name = get_valid_input(
                f"Enter citation format ({'/'.join(format_choices)}): ", 
                lambda x: x.lower() in format_choices,
                "apa"
            )
            
            id_list_option = input("Do you have a file with identifiers for the bibliography? (y/n): ")
            id_list_path = None
            id_type = "auto"
            
            if id_list_option.lower() == "y":
                id_list_path = input("Enter identifier list file path: ")
                if not os.path.exists(id_list_path):
                    print(f"File not found: {id_list_path}")
                    continue
                
                id_type = get_valid_input(
                    f"Enter ID type ({'/'.join(id_type_choices)}): ", 
                    lambda x: x.lower() in id_type_choices,
                    "auto"
                )
            
            print("\nProcessing Word document...")
            
            # Create args object for the handle_word_command function
            class Args:
                pass
            
            args = Args()
            args.input = input_path
            args.output = output_path
            args.format = format_name
            args.id_list = id_list_path
            args.id_type = id_type
            
            handle_word_command(args, citation_manager)
        
        elif choice == "5":
            # Exit
            print("\nExiting Citation Manager. Goodbye!")
            break
        
        else:
            print("\nInvalid choice. Please enter a number between 1 and 5.")

def get_valid_input(prompt, validator, default=None):
    """
    Get valid input from the user
    
    Args:
        prompt: Prompt to display
        validator: Function to validate input
        default: Default value if input is empty
    
    Returns:
        Valid input
    """
    while True:
        user_input = input(prompt)
        
        if not user_input and default:
            return default
        
        if validator(user_input):
            return user_input.lower()
        
        print("Invalid input. Please try again.")

if __name__ == "__main__":
    main()
