import requests
import json
import sys
import os
import re
import argparse
from typing import Dict, Any, List, Optional, Tuple
import xml.etree.ElementTree as ET
from datetime import datetime
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt


class NumberedCitationProcessor:
    """Process citations with a numbered system like [1], [2], etc."""
    
    def __init__(self, citation_manager):
        """Initialize with citation manager"""
        self.citation_manager = citation_manager
        self.citations_order = {}  # Maps metadata keys to citation numbers
        self.current_number = 1
    
    def process_document(self, document, citations, extracted_metadata, format_name):
        """
        First pass to collect and number all citations
        
        Args:
            document: Word document
            citations: Dictionary of extracted citation references
            extracted_metadata: Metadata dictionary
            format_name: Citation format name
            
        Returns:
            Dictionary mapping citation keys to citation numbers
        """
        # Reset numbering
        self.citations_order = {}
        self.current_number = 1
        
        # First, get all unique citations
        for paragraph in document.paragraphs:
            self._collect_citations_from_text(paragraph.text, citations)
            
        # Then for tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._collect_citations_from_text(paragraph.text, citations)
        
        print(f"Found {len(self.citations_order)} unique citations for numbering")
        return self.citations_order
    
    def _collect_citations_from_text(self, text, citations):
        """Collect citations from text and assign numbers"""
        for citation_key, citation in citations.items():
            if "source_text" in citation and citation["source_text"] in text:
                # Check if we have metadata for this citation
                if "metadata_key" in citation and citation["metadata_key"] not in self.citations_order:
                    self.citations_order[citation["metadata_key"]] = self.current_number
                    self.current_number += 1
    
    def format_in_text_citations(self, document, citations, citation_numbers):
        """
        Replace in-text citations with numbered format [1], [2], etc.
        
        Args:
            document: Word document
            citations: Dictionary of citation references
            citation_numbers: Dictionary mapping metadata keys to numbers
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Process paragraphs
            for paragraph in document.paragraphs:
                self._update_paragraph_with_numbers(paragraph, citations, citation_numbers)
            
            # Process tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._update_paragraph_with_numbers(paragraph, citations, citation_numbers)
            
            return True
        except Exception as e:
            print(f"Error formatting numbered citations: {e}")
            return False
    
    def _update_paragraph_with_numbers(self, paragraph, citations, citation_numbers):
        """Update paragraph with numbered citations"""
        text = paragraph.text
        updated_text = text
        
        # Replace each citation with its number
        for citation_key, citation in citations.items():
            if "source_text" in citation and citation["source_text"] in text:
                if "metadata_key" in citation and citation["metadata_key"] in citation_numbers:
                    number = citation_numbers[citation["metadata_key"]]
                    updated_text = updated_text.replace(
                        citation["source_text"],
                        f"[{number}]"
                    )
        
        # Update the paragraph text if changes were made
        if updated_text != text:
            paragraph.text = updated_text
    
    def generate_numbered_bibliography(self, document, extracted_metadata, citation_numbers, format_name):
        """
        Generate a numbered bibliography in the document
        
        Args:
            document: Word document
            extracted_metadata: Dictionary of extracted metadata
            citation_numbers: Dictionary mapping metadata keys to numbers
            format_name: Citation format name
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Add bibliography section
            document.add_page_break()
            try:
                document.add_heading("References", level=1)
            except Exception as e:
                print(f"Warning: Could not add styled heading: {e}")
                # Alternative: just add a paragraph with "References"
                p = document.add_paragraph("References")
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(16)  # Approximation of heading size
            
            # Get formatter
            formatter = self.citation_manager.formatters.get(format_name.lower())
            if not formatter:
                print(f"Error: Formatter '{format_name}' not found")
                return False
            
            # Create ordered list of citations
            ordered_citations = []
            for metadata_key, number in citation_numbers.items():
                if metadata_key in extracted_metadata:
                    metadata = extracted_metadata[metadata_key]
                    bibliography = formatter.format_citation(metadata)
                    ordered_citations.append((number, bibliography))
            
            # Sort by citation number
            ordered_citations.sort()
            
            # Add each citation to the bibliography
            for number, bibliography in ordered_citations:
                paragraph = document.add_paragraph()
                paragraph.text = f"{number}. {bibliography}"
            
            return True
        except Exception as e:
            print(f"Error generating numbered bibliography: {e}")
            return False