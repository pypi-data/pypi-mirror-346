# ===== Word Document Processing Classes =====
import requests
import json
import sys
import os
import re
import argparse
from typing import Dict, Any, List, Optional, Tuple, Union
import xml.etree.ElementTree as ET
from datetime import datetime
from urllib.parse import urlparse
from docx import Document
class CommandLineWordProcessor:
    """Process Word documents for citation management in command-line mode"""
    
    def __init__(self, citation_manager):
        """
        Initialize the word processor
        
        Args:
            citation_manager: The citation manager instance
        """
        self.citation_manager = citation_manager
        self.citation_pattern = r"\(\s*([^)]+)\s*,\s*(\d{4})\s*\)"  # Basic pattern for (Author, Year)
        
        # Additional patterns for different citation styles
        self.citation_patterns = {
            "author_year": r"\(\s*([^)]+)\s*,\s*(\d{4})\s*\)",  # (Author, Year)
            "superscript": r"(\w+)\s*(\d+)",  # Word1 for superscript citations
            "numbered": r"\[(\d+)\]",  # [1] for numbered citations
            "author_only": r"([A-Z][a-z]+)\s+et\s+al\.",  # Smith et al. for author only citations
        }
    
    def process_document(self, file_path: str, format_name: str, id_detector) -> Tuple[Document, Dict[str, Dict[str, Any]]]:
        """
        Process a Word document to extract citations
        
        Args:
            file_path: Path to the Word document
            format_name: Citation format name
            id_detector: ID detector instance
            
        Returns:
            Tuple of Document object and dictionary of citation metadata
        """
        print(f"Processing document: {file_path}")
        document = Document(file_path)
        citations = {}
        
        # Extract citations from paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text
            self._process_text_for_citations(text, citations, id_detector)
        
        # Extract citations from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_text_for_citations(paragraph.text, citations, id_detector)
        
        print(f"Found {len(citations)} potential citations")
        return document, citations
    
    def _process_text_for_citations(self, text: str, citations: Dict[str, Dict[str, Any]], id_detector) -> None:
        """
        Process text to extract citations
        
        Args:
            text: Text to process
            citations: Dictionary to store citations
            id_detector: ID detector instance
        """
        # Try all citation patterns
        for pattern_name, pattern in self.citation_patterns.items():
            matches = re.findall(pattern, text)
            for match in matches:
                if pattern_name == "author_year":
                    author, year = match
                    citation_key = f"{author.strip()}-{year.strip()}"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "author": author.strip(),
                            "year": year.strip(),
                            "pattern": pattern_name,
                            "source_text": f"({author.strip()}, {year.strip()})"
                        }
                elif pattern_name == "numbered":
                    number = match
                    citation_key = f"ref-{number}"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "number": number,
                            "pattern": pattern_name,
                            "source_text": f"[{number}]"
                        }
                elif pattern_name == "author_only":
                    author = match
                    citation_key = f"{author.strip()}-et-al"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "author": author.strip(),
                            "pattern": pattern_name,
                            "source_text": f"{author.strip()} et al."
                        }
        
        # Look for DOIs, PMIDs, arXiv IDs, and URLs in the text
        words = text.split()
        for word in words:
            word = word.strip(".,;()[]{}\"'")
            if word:
                id_type = id_detector.detect_id_type(word)
                if id_type != "unknown":
                    citation_key = f"{id_type}-{word}"
                    if citation_key not in citations:
                        citations[citation_key] = {
                            "id_type": id_type,
                            "id_value": word,
                            "pattern": "direct_id",
                            "source_text": word
                        }
    
    def update_document_with_citations(self, document: Document, 
                                  citations: Dict[str, Dict[str, Any]],
                                  extracted_metadata: Dict[str, Dict[str, Any]],
                                  format_name: str,
                                  output_path: str) -> bool:
        """
        Update document with formatted citations and add bibliography
        
        Args:
            document: The Word document object
            citations: Dictionary of extracted citation references
            extracted_metadata: Dictionary of extracted metadata
            format_name: Citation format name
            output_path: Path to save the updated document
            
        Returns:
            True if successful, False otherwise
        """
        # Create a mapping from citation keys to formatted citations
        formatted_citations = {}
        try:
            print(f"Updating document with {len(extracted_metadata)} citations in {format_name} format")
            
            # Get the formatter
            formatter = self.citation_manager.formatters.get(format_name.lower())
            if not formatter:
                print(f"Error: Formatter '{format_name}' not found")
                return False
            
            # Format each citation
            for citation_key, metadata in extracted_metadata.items():
                try:
                    formatted_in_text = formatter.format_in_text_citation(metadata)
                    formatted_citations[citation_key] = {
                        "in_text": formatted_in_text,
                        "bibliography": formatter.format_citation(metadata)
                    }
                    print(f"Formatted citation: {formatted_in_text} -> {formatter.format_citation(metadata)[:80]}...")
                except Exception as e:
                    print(f"Warning: Could not format citation {citation_key}: {e}")
                    continue
            
            # Update in-text citations - enhanced implementation
            try:
                print("Updating in-text citations...")
                citation_count = 0
                for paragraph in document.paragraphs:
                    changed = self._update_paragraph_citations(paragraph, citations, formatted_citations)
                    if changed:
                        citation_count += 1
                
                # Update citations in tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                changed = self._update_paragraph_citations(paragraph, citations, formatted_citations)
                                if changed:
                                    citation_count += 1
                                    
                print(f"Updated {citation_count} in-text citations")
            except Exception as e:
                print(f"Warning: Error updating in-text citations: {e}")
                # Continue processing even if in-text citations fail
        except Exception as e:
            print(f"Warning: Error in citation formatting: {e}")
            # Continue processing even if formatting fails
        
        # Create list of citations in proper order
        ordered_citations = []
        try:
            # Get unique metadata keys from citations
            used_metadata_keys = set()
            for citation in citations.values():
                if "metadata_key" in citation and citation["metadata_key"] in extracted_metadata:
                    used_metadata_keys.add(citation["metadata_key"])
            
            # Add any metadata keys that weren't explicitly linked to citations
            for key in extracted_metadata.keys():
                used_metadata_keys.add(key)
            
            # Create ordered list of bibliography entries
            ordered_citations = []
            for key in used_metadata_keys:
                if key in formatted_citations:
                    ordered_citations.append((key, formatted_citations[key]["bibliography"]))
            
            # Sort by first author's last name and year
            def get_sort_key(citation_tuple):
                key, bibliography = citation_tuple
                metadata = extracted_metadata.get(key, {})
                authors = metadata.get("authors", [])
                year = metadata.get("year", "9999")  # Default to high year if not found
                
                if authors:
                    # Get first author's last name
                    author = authors[0]
                    if "," in author:
                        last_name = author.split(",")[0].lower()
                    else:
                        parts = author.split()
                        last_name = parts[-1].lower() if parts else ""
                    return (last_name, year)
                return ("zzzzzz", year)  # Sort unknown authors at end
            
            # Sort the citations
            ordered_citations.sort(key=get_sort_key)
            
            print(f"Prepared {len(ordered_citations)} bibliography entries in alphabetical order")
        except Exception as e:
            print(f"Warning: Error organizing bibliography entries: {e}")
            # Fall back to using all available formatted citations
            ordered_citations = [(k, v["bibliography"]) for k, v in formatted_citations.items()]
            
        try:
            # Add bibliography section
            try:
                document.add_page_break()
            except Exception as e:
                print(f"Warning: Could not add page break: {e}")
                # Continue without page break
                
            # Create References title - with multiple fallback methods
            try:
                # Method 1: Try to add heading with style
                document.add_heading("References", level=1)
            except Exception as e:
                try:
                    # Method 2: Try to add a plain paragraph with bold text
                    print(f"Using simple heading instead of styled heading: {e}")
                    paragraph = document.add_paragraph("References")
                    paragraph.style = 'Normal'  # This will always exist
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.size = Pt(14)  # Approximation of heading size
                    
                    # Try to set font size if possible
                    try:
                        from docx.shared import Pt
                        run.font.size = Pt(16)  # Approximate size for Heading 1
                    except ImportError:
                        # If docx.shared is not available, just use bold
                        pass
                except Exception as e2:
                    # Method 3: Last resort - just add plain text
                    print(f"Warning: Could not add bold heading: {e2}")
                    document.add_paragraph("REFERENCES")
                    
            # Add each citation to the bibliography
            if ordered_citations:
                print(f"Adding {len(ordered_citations)} citations to bibliography")
                for i, (citation_key, bibliography) in enumerate(ordered_citations, 1):
                    try:
                        paragraph = document.add_paragraph()
                        # Add number for numbered citations
                        paragraph.text = f"{i}. {bibliography}"
                        print(f"Added bibliography entry {i}: {bibliography[:80]}...")
                    except Exception as e:
                        print(f"Warning: Could not add bibliography entry for {citation_key}: {e}")
                        continue
            else:
                print("No citations to add to bibliography")
        except Exception as e:
            print(f"Warning: Error adding bibliography: {e}")
            # Continue without bibliography if it fails
        
        # Save the document - this is the most critical part
        try:
            document.save(output_path)
            print(f"Document saved to {output_path}")
            return True
        except Exception as e:
            print(f"Error saving document to {output_path}: {e}")
            
            # Try an alternative approach to saving if the first one fails
            try:
                # Create a new document and copy content
                from docx import Document as NewDocument
                new_doc = NewDocument()
                
                # Copy the paragraphs from the original document
                for para in document.paragraphs:
                    new_para = new_doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                
                # Save the new document
                new_doc.save(output_path)
                print(f"Document saved using alternative method to {output_path}")
                return True
            except Exception as e2:
                print(f"Fatal error: Could not save document using alternative method: {e2}")
                return False
    
    def _update_paragraph_citations(self, paragraph, citations, formatted_citations) -> bool:
        """
        Update citations in a paragraph
        
        Args:
            paragraph: The paragraph object
            citations: Dictionary of citation references
            formatted_citations: Dictionary of formatted citations
        
        Returns:
            True if paragraph was modified, False otherwise
        """
        text = paragraph.text
        updated_text = text
        modified = False
        
        # Replace each citation with the formatted version
        for citation_key, citation in citations.items():
            if "source_text" in citation and citation["source_text"] in text:
                # Find the metadata key for this citation
                if "metadata_key" in citation and citation["metadata_key"] in formatted_citations:
                    metadata_key = citation["metadata_key"]
                    formatted_citation = formatted_citations[metadata_key]["in_text"]
                    
                    # Only replace if we have something valid
                    if formatted_citation:
                        updated_text = updated_text.replace(
                            citation["source_text"], 
                            formatted_citation
                        )
                        modified = True
                else:
                    # Try to find a direct match in formatted_citations
                    for fmt_key, fmt_citation in formatted_citations.items():
                        if citation_key in fmt_key:
                            updated_text = updated_text.replace(
                                citation["source_text"],
                                fmt_citation["in_text"]
                            )
                            modified = True
                            break
        
        # Update the paragraph text if changes were made
        if updated_text != text:
            paragraph.text = updated_text
            return True
        
        return False
    
    def _get_matching_key(self, citation, formatted_citations) -> Optional[str]:
        """
        Find a matching key in the formatted citations
        
        Args:
            citation: Citation reference
            formatted_citations: Dictionary of formatted citations
            
        Returns:
            Matching key if found, None otherwise
        """
        # This is a simplified matching logic
        # A more sophisticated implementation would use fuzzy matching
        
        # Direct match by ID
        if "id_type" in citation and "id_value" in citation:
            for key in formatted_citations:
                if citation["id_type"] in key and citation["id_value"] in key:
                    return key
        
        # Match by author and year
        if "author" in citation and "year" in citation:
            author = citation["author"].lower()
            year = citation["year"]
            
            for key in formatted_citations:
                metadata_key = key.split("-metadata-")[0] if "-metadata-" in key else key
                if author in metadata_key.lower() and year in metadata_key:
                    return key
        
        return None