from typing import Any
import httpx
import os
import re
import subprocess
import tempfile
import uuid
from pathlib import Path
import requests
import json
import base64
from io import BytesIO
from bs4 import BeautifulSoup
import time
import datetime
from colorama import init, Fore, Back, Style
import math
from PIL import Image
import sys
import inspect

# 打印当前Python路径和版本，帮助诊断环境问题
print(f"当前Python路径: {sys.executable}")
print(f"当前Python版本: {sys.version}")
print(f"当前系统路径: {sys.path}")

import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 初始化colorama
init()

# 定义调试信息打印函数
def debug_print(message, level="INFO", indent=0, show_time=True):
    """
    格式化打印调试信息
    
    Args:
        message: 要打印的消息
        level: 消息级别 (INFO, WARNING, ERROR, SUCCESS, DEBUG)
        indent: 缩进级别
        show_time: 是否显示时间戳
    """
    # 颜色映射
    colors = {
        "INFO": Fore.CYAN,
        "WARNING": Fore.YELLOW,
        "ERROR": Fore.RED,
        "SUCCESS": Fore.GREEN,
        "DEBUG": Fore.MAGENTA,
        "HEADER": Fore.BLUE + Style.BRIGHT
    }
    
    # 获取当前时间
    timestamp = ""
    if show_time:
        now = datetime.datetime.now()
        timestamp = f"[{now.strftime('%H:%M:%S')}] "
    
    # 构建缩进
    indent_str = "  " * indent
    
    # 构建前缀
    prefix = f"{colors.get(level, '')}{timestamp}{level.ljust(7)}{Style.RESET_ALL} "
    
    # 打印消息
    print(f"{prefix}{indent_str}{message}")

def print_section_header(title, width=80):
    """打印带有分隔线的章节标题"""
    debug_print("", show_time=False)
    debug_print("=" * width, level="HEADER", show_time=False)
    debug_print(f"= {title.center(width-4)} =", level="HEADER", show_time=False)
    debug_print("=" * width, level="HEADER", show_time=False)
    debug_print("", show_time=False)

def print_section_footer(title, width=80):
    """打印带有分隔线的章节结束标记"""
    debug_print("", show_time=False)
    debug_print("-" * width, level="HEADER", show_time=False)
    debug_print(f"- {('END: ' + title).center(width-4)} -", level="HEADER", show_time=False)
    debug_print("-" * width, level="HEADER", show_time=False)
    debug_print("", show_time=False)

def print_progress(current, total, prefix="", suffix="", length=50):
    """打印进度条"""
    percent = float(current) / float(total)
    filled_length = int(length * percent)
    bar = "█" * filled_length + "░" * (length - filled_length)
    progress_str = f"{prefix} |{bar}| {current}/{total} {suffix} ({percent:.1%})"
    debug_print(progress_str, level="INFO", show_time=False)

# 定义多个主题配置
MERMAID_THEMES = {
    'default': """%%{
  init: {
    'theme': 'default', 
    'themeVariables': {
      'fontSize': '18px',
      'fontFamily': 'Arial, sans-serif',
      'primaryColor': '#6C5CE7',
      'primaryTextColor': '#fff',
      'primaryBorderColor': '#4834D4',
      'lineColor': '#A55EEA',
      'secondaryColor': '#D6A2E8',
      'tertiaryColor': '#fff',
      'edgeLabelBackground': '#ffffff',
      'nodeTextColor': '#000',
      'mainBkg': '#DFE6E9',
      'secondBkg': '#F8F9FA',
      'clusterBkg': '#fdffbc',
      'titleColor': '#2D3436'
    }
  }
}%%
""",
    
    'dark': """%%{
  init: {
    'theme': 'dark',
    'themeVariables': {
      'fontSize': '16px',
      'fontFamily': 'Menlo, monospace',
      'primaryColor': '#FF79C6',
      'primaryTextColor': '#282A36',
      'primaryBorderColor': '#BD93F9',
      'lineColor': '#FF79C6',
      'secondaryColor': '#6272A4',
      'tertiaryColor': '#44475A',
      'edgeLabelBackground': '#282A36',
      'nodeTextColor': '#F8F8F2',
      'mainBkg': '#282A36',
      'secondBkg': '#44475A',
      'clusterBkg': '#44475A',
      'titleColor': '#F8F8F2'
    }
  }
}%%
""",

    'forest': """%%{
  init: {
    'theme': 'forest',
    'themeVariables': {
      'fontSize': '16px',
      'fontFamily': 'Georgia, serif',
      'primaryColor': '#2ECC71',
      'primaryTextColor': '#FFF',
      'primaryBorderColor': '#27AE60',
      'lineColor': '#27AE60',
      'secondaryColor': '#95A5A6',
      'tertiaryColor': '#ECF0F1',
      'edgeLabelBackground': '#FFF',
      'nodeTextColor': '#2C3E50',
      'mainBkg': '#F1F8E9',
      'secondBkg': '#DCEDC8',
      'clusterBkg': '#C8E6C9',
      'titleColor': '#1B5E20'
    }
  }
}%%
""",

    'ocean': """%%{
  init: {
    'theme': 'ocean',
    'themeVariables': {
      'fontSize': '16px',
      'fontFamily': 'Helvetica, sans-serif',
      'primaryColor': '#3498DB',
      'primaryTextColor': '#FFF',
      'primaryBorderColor': '#2980B9',
      'lineColor': '#2980B9',
      'secondaryColor': '#E74C3C',
      'tertiaryColor': '#ECF0F1',
      'edgeLabelBackground': '#FFF',
      'nodeTextColor': '#2C3E50',
      'mainBkg': '#EBF5FB',
      'secondBkg': '#D6EAF8',
      'clusterBkg': '#D4E6F1',
      'titleColor': '#21618C'
    }
  }
}%%
""",

    'elegant': """%%{
  init: {
    'theme': 'elegant',
    'themeVariables': {
      'fontSize': '17px',
      'fontFamily': 'Palatino, serif',
      'primaryColor': '#34495E',
      'primaryTextColor': '#FFF',
      'primaryBorderColor': '#2C3E50',
      'lineColor': '#2C3E50',
      'secondaryColor': '#95A5A6',
      'tertiaryColor': '#ECF0F1',
      'edgeLabelBackground': '#FFF',
      'nodeTextColor': '#2C3E50',
      'mainBkg': '#FDFEFE',
      'secondBkg': '#F8F9F9',
      'clusterBkg': '#F4F6F6',
      'titleColor': '#2C3E50'
    }
  }
}%%
"""
}

def beautify_mermaid_code(mermaid_code, theme='default'):
    """
    美化Mermaid代码，使生成的图表更美观
    
    Args:
        mermaid_code: 原始的Mermaid代码
        theme: 主题名称，可选值：'default', 'dark', 'forest', 'ocean', 'elegant'
        
    Returns:
        美化后的Mermaid代码
    """
    # 检测图表类型
    lines = mermaid_code.strip().split('\n')
    if not lines:
        return mermaid_code
    
    # 获取选定的主题配置
    theme_config = MERMAID_THEMES.get(theme, MERMAID_THEMES['default'])
    
    # 为所有图表应用主题配置，除非已经有配置
    if not mermaid_code.strip().startswith("%%{"):
        mermaid_code = theme_config + mermaid_code
    
    return mermaid_code

def render_mermaid_to_image(mermaid_code, output_path=None, theme='default'):
    """
    使用Kroki API将Mermaid图表渲染为图像。
    返回保存的图像路径。
    
    Args:
        mermaid_code: Mermaid图表代码
        output_path: 可选输出文件路径，默认创建临时文件
        theme: 主题名称
        
    Returns:
        图像文件路径，渲染失败则返回None
    """
    # 创建临时文件保存图像（如果未提供输出路径）
    if not output_path:
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, f"mermaid_{uuid.uuid4()}.png")
    
    # 美化Mermaid代码
    mermaid_code = beautify_mermaid_code(mermaid_code, theme)
    
    # 方法1: 尝试使用mermaid-py库（如果已安装）
    # try:
    #     debug_print("尝试使用mermaid-py库...", level="INFO")
    #     import importlib
        
    #     # 检查mermaid模块是否已安装
    #     if importlib.util.find_spec("mermaid"):
    #         debug_print("找到mermaid模块", level="INFO")
    #         import mermaid as md
    #         from mermaid.graph import Graph
            
    #         debug_print("使用mermaid-py生成图像...", level="INFO")
            
    #         # 创建图表对象 - 按照官方文档示例
    #         graph = Graph('mermaid-diagram', mermaid_code)
    #         render = md.Mermaid(graph)
            
    #         # 设置样式和大小 - 使用对象的属性和方法
    #         # 注意: 这些是私有属性，但从用户提供的方法列表可以看出它们是可用的
    #         debug_print("设置Mermaid图表样式和大小...", level="INFO")
            
    #         # 设置图表宽度和高度 (如果有这些属性)
    #         if hasattr(render, '_Mermaid__width'):
    #             render._Mermaid__width = 800  # 设置适当宽度
    #         if hasattr(render, '_Mermaid__height'):
    #             render._Mermaid__height = 600  # 设置适当高度
            
    #         # 设置缩放因子 (会影响字体大小)
    #         if hasattr(render, '_Mermaid__scale'):
    #             render._Mermaid__scale = 1.2  # 设置比默认大一点的缩放比例
            
    #         # 如果有set_position方法，尝试使用它
    #         if hasattr(render, 'set_position'):
    #             try:
    #                 render.set_position(position="center")  # 居中显示
    #             except Exception as position_error:
    #                 debug_print(f"设置position失败: {position_error}", level="WARNING")
            
    #         # 尝试保存为图像
    #         try:
    #             # 查看现有方法
    #             debug_print(f"可用方法: {dir(render)}", level="DEBUG")
                
    #             # 优先使用to_png方法(如果存在)，因为用户提供的方法列表显示这是可用的
    #             if hasattr(render, 'to_png'):
    #                 debug_print("使用to_png方法", level="INFO")
    #                 render.to_png(output_path)
    #             # 最新版本尝试使用render_image方法
    #             elif hasattr(render, 'render_image'):
    #                 debug_print("使用render_image方法", level="INFO")
    #                 with open(output_path, 'wb') as f:
    #                     f.write(render.render_image())
    #             # 尝试使用save_as_image方法
    #             elif hasattr(render, 'save_as_image'):
    #                 debug_print("使用save_as_image方法", level="INFO")
    #                 render.save_as_image(output_path)
    #             # 如果不存在，可能需要其他方法
    #             elif hasattr(render, 'export'):
    #                 debug_print("使用export方法", level="INFO")
    #                 render.export(output_path, 'png')
    #             else:
    #                 debug_print("无法找到适当的图像导出方法", level="WARNING")
    #                 raise AttributeError("mermaid对象没有支持的图像导出方法")
                
    #             # 验证图像是否创建成功
    #             if os.path.exists(output_path):
    #                 debug_print("mermaid-py渲染成功!", level="SUCCESS")
    #                 return output_path
    #         except Exception as e:
    #             debug_print(f"mermaid-py导出图像失败: {e}", level="WARNING")
    #     else:
    #         debug_print("mermaid-py库未安装，尝试其他方法", level="INFO")
    # except Exception as e:
    #     debug_print(f"使用mermaid-py出错: {e}", level="WARNING")
    
    # 方法2: 使用 Kroki API 渲染图表
    for retry in range(2):  # 添加重试逻辑
        try:
            print(f"尝试使用 Kroki API... (尝试 {retry+1}/2)")
            payload = {
                "diagram_source": mermaid_code,
                "diagram_type": "mermaid",
                "output_format": "png"
            }
            
            print("发送请求到 Kroki API...")
            response = requests.post("https://kroki.io/mermaid/png", json=payload, timeout=15)
            
            if response.status_code == 200:
                # 保存图像
                with open(output_path, 'wb') as f:
                    f.write(response.content)
                
                # 检查生成的图片是否有效
                try:
                    img = Image.open(output_path)
                    width, height = img.size
                    if width < 10 or height < 10:  # 检查图片是否太小
                        print(f"生成的图片太小 ({width}x{height})，可能是渲染失败")
                        if retry < 1:
                            continue
                    img.close()
                except Exception as img_error:
                    print(f"检查图片时出错: {img_error}")
                    if retry < 1:
                        continue
                
                # 后处理图像 - 缩放来减小字体大小
                process_image_scaling(output_path, scale_factor=0.9)
                print("Kroki API 渲染成功!")
                return output_path
            else:
                print(f"Kroki API请求失败，状态码: {response.status_code}")
                print(f"响应内容: {response.text[:100] if hasattr(response, 'text') else 'No text response'}")
                if retry < 1:
                    print("将重试...")
                    time.sleep(1)
        except Exception as e:
            print(f"使用 Kroki API 出错: {e}")
            if retry < 1:
                print("将重试...")
                time.sleep(1)
    
    # 渲染失败 - 生成一个简单的错误图像
    try:
        print("所有渲染方法都失败，生成简单的错误图像...")
        # 创建一个包含错误信息的图像
        img = Image.new("RGB", (600, 300), color=(255, 255, 255))
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        
        # 尝试使用系统字体，如果失败则使用默认字体
        try:
            font = ImageFont.truetype("Arial", 14)
        except:
            font = ImageFont.load_default()
        
        # 绘制错误信息
        draw.text((20, 20), "Mermaid 图表渲染失败", fill=(255, 0, 0), font=font)
        draw.text((20, 50), "请检查:", fill=(0, 0, 0), font=font)
        draw.text((30, 80), "1. Mermaid 语法是否正确", fill=(0, 0, 0), font=font)
        draw.text((30, 110), "2. 网络连接是否正常", fill=(0, 0, 0), font=font)
        draw.text((30, 140), "3. 安装 mermaid-py==0.7.0 以提高渲染成功率", fill=(0, 0, 0), font=font)
        
        # 显示部分图表代码
        code_preview = mermaid_code[:200] + "..." if len(mermaid_code) > 200 else mermaid_code
        code_lines = code_preview.split("\n")
        for i, line in enumerate(code_lines[:5]):
            if i >= 5:
                break
            draw.text((20, 180 + i * 20), line, fill=(0, 0, 255), font=font)
        
        img.save(output_path)
        print("已生成错误信息图像")
        return output_path
    except Exception as e:
        print(f"生成错误图像失败: {e}")
    
    print("无法生成图像")
    return None

def process_image_scaling(image_path, scale_factor=0.9):
    """对图像进行缩放处理，缩小字体大小"""
    try:
        # 打开图像文件
        img = Image.open(image_path)
        
        # 获取原始尺寸
        width, height = img.size
        print(f"原始图像尺寸: {width}x{height}")
        
        # 计算新尺寸
        new_width = math.ceil(width * scale_factor)
        new_height = math.ceil(height * scale_factor)
        print(f"缩放后图像尺寸: {new_width}x{new_height}")
        
        # 缩放图像
        resized_img = img.resize((new_width, new_height), Image.LANCZOS)
        
        # 保存处理后的图像
        resized_img.save(image_path)
        print(f"图像缩放完成: {image_path}")
        
        return True
    except Exception as e:
        print(f"图像处理失败: {e}")
        return False

def extract_mermaid_blocks(md_content):
    """Extract Mermaid code blocks from Markdown content."""
    # Pattern to match ```mermaid ... ``` blocks
    pattern = r'```mermaid\s+(.*?)\s+```'
    # Find all matches using re.DOTALL to match across multiple lines
    matches = re.findall(pattern, md_content, re.DOTALL)
    return matches


def extract_code_blocks(md_content):
    """Extract all code blocks (except Mermaid) from Markdown content."""
    # 使用更精确的模式匹配代码块
    # 匹配 ```language 和 ``` 之间的内容
    pattern = r'```(?!mermaid)([^\n]*)\n([\s\S]*?)\n```'
    matches = re.findall(pattern, md_content, re.DOTALL)
    
    # 处理匹配结果，清理语言和代码内容
    cleaned_matches = []
    for lang, code in matches:
        # 清理语言标识
        lang = lang.strip()
        # 清理代码内容，但保留换行符
        code = code.rstrip()  # 只去除尾部空白，保留换行
        cleaned_matches.append((lang, code))
    
    print(f"Found {len(cleaned_matches)} code blocks")
    for i, (lang, code) in enumerate(cleaned_matches):
        print(f"Code block {i}: language='{lang}', length={len(code)} chars")
        if len(code) > 50:
            preview = code.split('\n')[0][:50]  # 只显示第一行的前50个字符
            print(f"  Preview (first line): {preview}...")
        else:
            print(f"  Content: {code}")
    
    return cleaned_matches


def html_to_docx(html_content, doc, table_style='Table Grid'):
    """Convert HTML content to Word document elements."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 首先检查是否有代码块元素
    code_elements = soup.find_all('code')
    if code_elements:
        print(f"Found {len(code_elements)} code elements in HTML")
    
    pre_elements = soup.find_all('pre')
    if pre_elements:
        print(f"Found {len(pre_elements)} pre elements in HTML")
        # 处理pre元素（通常包含代码块）
        for pre in pre_elements:
            code = pre.get_text(strip=True)
            if code:
                print(f"Processing pre element with content length: {len(code)}")
                # 检查是否有语言类
                lang = ""
                if pre.has_attr('class'):
                    for cls in pre['class']:
                        if cls.startswith('language-'):
                            lang = cls.replace('language-', '')
                            break
                
                # 使用代码块格式化函数
                format_code_block(doc, code, lang)
                # 从soup中移除已处理的元素，避免重复处理
                pre.extract()
    
    # Process elements in order
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote', 'table']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            paragraph.style = f'Heading {element.name[1]}'
        
        elif element.name == 'p':
            # 检查段落中是否包含代码元素
            code_in_p = element.find('code')
            if code_in_p and len(element.contents) == 1:
                # 如果段落只包含一个代码元素，将其作为代码块处理
                code = code_in_p.get_text(strip=True)
                lang = ""
                if code_in_p.has_attr('class'):
                    for cls in code_in_p['class']:
                        if cls.startswith('language-'):
                            lang = cls.replace('language-', '')
                            break
                format_code_block(doc, code, lang)
            else:
                # 正常处理段落，不添加边框
                paragraph = doc.add_paragraph(element.get_text(strip=True))
                apply_style_to_paragraph(paragraph, element)
        
        elif element.name == 'blockquote':
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            paragraph.style = 'Quote'
        
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                process_list_item(doc, li, 'List Bullet')
        
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                process_list_item(doc, li, 'List Number')
        
        elif element.name == 'table':
            process_table(doc, element, table_style)
    
    return doc


def apply_style_to_paragraph(paragraph, element):
    """Apply HTML styles to a Word paragraph based on the element."""
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        paragraph.style = f'Heading {level}'
    
    if element.name == 'strong' or element.find('strong'):
        for run in paragraph.runs:
            run.bold = True
    
    if element.name == 'em' or element.find('em'):
        for run in paragraph.runs:
            run.italic = True
    
    if element.name == 'u' or element.find('u'):
        for run in paragraph.runs:
            run.underline = True
    
    if element.name == 'code' or element.find('code'):
        for run in paragraph.runs:
            run.font.name = 'Courier New'
    
    if element.name == 'center' or element.get('align') == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def process_list_item(doc, li_element, list_style, level=0):
    """Process a list item and its children with proper indentation."""
    # Add the list item with proper style and level
    text = li_element.get_text(strip=True)
    paragraph = doc.add_paragraph(text)
    paragraph.style = list_style
    paragraph.paragraph_format.left_indent = Pt(18 * level)  # Indent based on nesting level
    
    # Process any nested lists
    nested_ul = li_element.find('ul')
    nested_ol = li_element.find('ol')
    
    if nested_ul:
        for nested_li in nested_ul.find_all('li', recursive=False):
            process_list_item(doc, nested_li, 'List Bullet', level + 1)
    
    if nested_ol:
        for nested_li in nested_ol.find_all('li', recursive=False):
            process_list_item(doc, nested_li, 'List Number', level + 1)


def process_table(doc, table_element, table_style='Table Grid'):
    """Process a table element and convert it to a Word table."""
    # Find all rows in the table
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count the maximum number of cells in any row
    max_cols = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        max_cols = max(max_cols, len(cells))
    
    if max_cols == 0:
        return
    
    # Create the table in the document
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = table_style
    
    # Fill the table with data
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < max_cols:  # Ensure we don't exceed the table dimensions
                # Get cell text and apply basic formatting
                text = cell.get_text(strip=True)
                table.cell(i, j).text = text
                
                # Apply header formatting if it's a header cell
                if cell.name == 'th' or i == 0:
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def add_border_to_paragraph(paragraph):
    """Add border to a paragraph."""
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # The order of these elements matters for Word
    for border_pos in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_pos}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width in 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        pBdr.append(border)
    
    # Insert the border element before these elements
    child_elements = [
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ]
    
    # Find the first child in the list that exists
    insert_after = None
    for child in child_elements:
        element = pPr.find(qn(child))
        if element is not None:
            insert_after = element
            break
    
    # If none of the specified children is found, just append to the end
    if insert_after is None:
        pPr.append(pBdr)
    else:
        insert_after.addprevious(pBdr)


def format_code_block(doc, code, language=""):
    """Format a code block with proper styling in the Word document."""
    print(f"Formatting code block with language: '{language}', code length: {len(code)} chars")
    
    # 创建一个段落用于代码块
    code_para = doc.add_paragraph()
    code_para.style = 'No Spacing'
    
    # 添加背景色
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')  # 浅灰色背景
    code_para._p.get_or_add_pPr().append(shading_elm)
    
    # 语言标识映射（规范化语言名称）
    language_map = {
        'js': 'javascript',
        'ts': 'typescript',
        'py': 'python',
        'cs': 'csharp',
        'c#': 'csharp',
        'rb': 'ruby',
        'go': 'golang',
        'rs': 'rust',
        'sh': 'bash',
        'yml': 'yaml',
        'md': 'markdown',
    }
    
    # 规范化语言名称
    lang_normalized = language.lower().strip()
    if lang_normalized in language_map:
        lang_normalized = language_map[lang_normalized]
        display_language = lang_normalized
    else:
        display_language = language.strip()
    
    # 如果有语言标识，添加到代码块前面
    if display_language:
        # 创建单独的段落用于语言标识
        lang_para = doc.add_paragraph()
        lang_para.style = 'No Spacing'
        
        # 添加语言标识背景
        lang_shading_elm = OxmlElement('w:shd')
        lang_shading_elm.set(qn('w:fill'), 'E0E0E0')  # 稍深的灰色背景
        lang_para._p.get_or_add_pPr().append(lang_shading_elm)
        
        # 添加语言文本
        lang_run = lang_para.add_run(f"{display_language}")
        lang_run.bold = True
        lang_run.font.size = Pt(9)
        lang_run.font.color.rgb = RGBColor(70, 70, 70)  # 深灰色文字
        
        # 为语言标识添加上边框和左右边框
        p = lang_para._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        for border_pos in ['top', 'left', 'right']:
            border = OxmlElement(f'w:{border_pos}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC')
            pBdr.append(border)
        
        pPr.append(pBdr)
    
    # 语言关键字定义
    language_keywords = {
        'javascript': [
            'function', 'return', 'if', 'else', 'for', 'while', 'do', 'switch', 'case', 'default',
            'break', 'continue', 'var', 'let', 'const', 'class', 'extends', 'new', 'this', 'super',
            'import', 'export', 'from', 'as', 'async', 'await', 'try', 'catch', 'finally', 'throw',
            'typeof', 'instanceof', 'in', 'of', 'delete', 'void', 'yield', 'static', 'get', 'set',
            'null', 'undefined', 'true', 'false'
        ],
        'typescript': [
            'function', 'return', 'if', 'else', 'for', 'while', 'do', 'switch', 'case', 'default',
            'break', 'continue', 'var', 'let', 'const', 'class', 'extends', 'new', 'this', 'super',
            'import', 'export', 'from', 'as', 'async', 'await', 'try', 'catch', 'finally', 'throw',
            'typeof', 'instanceof', 'in', 'of', 'delete', 'void', 'yield', 'static', 'get', 'set',
            'null', 'undefined', 'true', 'false', 'interface', 'type', 'enum', 'namespace',
            'readonly', 'private', 'protected', 'public', 'implements', 'declare', 'abstract',
            'any', 'boolean', 'number', 'string', 'symbol', 'unknown', 'never', 'void'
        ],
        'python': [
            'def', 'return', 'if', 'elif', 'else', 'for', 'while', 'break', 'continue', 'pass',
            'class', 'import', 'from', 'as', 'try', 'except', 'finally', 'raise', 'with',
            'assert', 'async', 'await', 'lambda', 'yield', 'global', 'nonlocal', 'in', 'is',
            'not', 'and', 'or', 'True', 'False', 'None', 'del', 'self'
        ],
        'java': [
            'public', 'private', 'protected', 'static', 'final', 'abstract', 'class', 'interface',
            'extends', 'implements', 'enum', 'new', 'this', 'super', 'return', 'if', 'else',
            'for', 'while', 'do', 'switch', 'case', 'default', 'break', 'continue', 'try', 'catch',
            'finally', 'throw', 'throws', 'instanceof', 'void', 'boolean', 'byte', 'char', 'short',
            'int', 'long', 'float', 'double', 'true', 'false', 'null', 'package', 'import'
        ],
        'csharp': [
            'public', 'private', 'protected', 'internal', 'static', 'readonly', 'const', 'virtual',
            'abstract', 'override', 'sealed', 'class', 'interface', 'struct', 'enum', 'namespace',
            'using', 'new', 'this', 'base', 'return', 'if', 'else', 'for', 'foreach', 'while', 'do',
            'switch', 'case', 'default', 'break', 'continue', 'try', 'catch', 'finally', 'throw',
            'void', 'bool', 'byte', 'sbyte', 'char', 'short', 'ushort', 'int', 'uint', 'long',
            'ulong', 'float', 'double', 'decimal', 'string', 'object', 'true', 'false', 'null',
            'var', 'dynamic', 'async', 'await', 'in', 'out', 'ref', 'params', 'where', 'is', 'as'
        ],
        'php': [
            'function', 'return', 'if', 'else', 'elseif', 'for', 'foreach', 'while', 'do', 'switch',
            'case', 'default', 'break', 'continue', 'class', 'interface', 'trait', 'extends',
            'implements', 'new', 'this', 'parent', 'public', 'private', 'protected', 'static',
            'abstract', 'final', 'try', 'catch', 'finally', 'throw', 'use', 'namespace', 'include',
            'include_once', 'require', 'require_once', 'echo', 'print', 'global', 'const', 'var',
            'true', 'false', 'null', 'array', 'list', 'isset', 'empty', 'unset'
        ],
        'ruby': [
            'def', 'end', 'if', 'else', 'elsif', 'unless', 'case', 'when', 'while', 'until', 'for',
            'break', 'next', 'redo', 'retry', 'return', 'class', 'module', 'include', 'extend',
            'attr_reader', 'attr_writer', 'attr_accessor', 'public', 'private', 'protected', 'begin',
            'rescue', 'ensure', 'raise', 'yield', 'super', 'self', 'nil', 'true', 'false', 'and',
            'or', 'not', 'alias', 'undef', 'defined?'
        ],
        'golang': [
            'func', 'return', 'if', 'else', 'for', 'range', 'switch', 'case', 'default', 'break',
            'continue', 'goto', 'fallthrough', 'package', 'import', 'type', 'struct', 'interface',
            'map', 'chan', 'go', 'select', 'defer', 'var', 'const', 'true', 'false', 'nil', 'make',
            'new', 'len', 'cap', 'append', 'delete', 'copy', 'close', 'panic', 'recover'
        ],
        'rust': [
            'fn', 'return', 'if', 'else', 'match', 'for', 'while', 'loop', 'break', 'continue',
            'let', 'mut', 'const', 'static', 'struct', 'enum', 'trait', 'impl', 'pub', 'mod',
            'use', 'crate', 'super', 'self', 'Self', 'where', 'async', 'await', 'move', 'dyn',
            'unsafe', 'extern', 'type', 'union', 'ref', 'box', 'in', 'true', 'false', 'as'
        ],
        'bash': [
            'if', 'then', 'else', 'elif', 'fi', 'case', 'esac', 'for', 'while', 'until', 'do',
            'done', 'in', 'function', 'time', 'select', 'break', 'continue', 'return', 'exit',
            'export', 'local', 'readonly', 'shift', 'source', 'alias', 'unalias', 'set', 'unset',
            'trap', 'declare', 'typeset', 'let', 'eval', 'exec', 'ulimit', 'umask', 'wait'
        ],
        'html': [
            'html', 'head', 'body', 'title', 'meta', 'link', 'script', 'style', 'div', 'span',
            'p', 'a', 'img', 'ul', 'ol', 'li', 'table', 'tr', 'td', 'th', 'form', 'input',
            'button', 'select', 'option', 'textarea', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
            'header', 'footer', 'nav', 'section', 'article', 'aside', 'main', 'canvas', 'audio',
            'video', 'source', 'iframe', 'strong', 'em', 'br', 'hr', 'pre', 'code', 'blockquote'
        ],
        'css': [
            '@media', '@import', '@keyframes', '@font-face', 'from', 'to', 'important',
            'margin', 'padding', 'border', 'width', 'height', 'color', 'background',
            'font-size', 'font-family', 'font-weight', 'text-align', 'display', 'position',
            'top', 'right', 'bottom', 'left', 'float', 'clear', 'overflow', 'z-index',
            'opacity', 'transition', 'transform', 'animation', 'flex', 'grid'
        ],
        'sql': [
            'select', 'from', 'where', 'and', 'or', 'not', 'in', 'between', 'like', 'is', 'null',
            'as', 'join', 'inner', 'left', 'right', 'outer', 'full', 'on', 'group', 'by', 'having',
            'order', 'asc', 'desc', 'limit', 'offset', 'union', 'all', 'insert', 'into', 'values',
            'update', 'set', 'delete', 'create', 'table', 'alter', 'drop', 'index', 'view',
            'procedure', 'function', 'trigger', 'primary', 'key', 'foreign', 'references',
            'constraint', 'default', 'int', 'varchar', 'text', 'date', 'datetime', 'timestamp',
            'boolean', 'float', 'double', 'decimal', 'auto_increment'
        ]
    }
    
    # 获取当前语言的关键字列表
    keywords = []
    if lang_normalized in language_keywords:
        keywords = language_keywords[lang_normalized]
    else:
        # 对于未知语言，使用通用关键字
        keywords = [
            'function', 'return', 'if', 'else', 'for', 'while', 'class', 'import', 
            'from', 'var', 'let', 'const', 'def', 'async', 'await', 'try', 'catch',
            'public', 'private', 'protected', 'static', 'void', 'int', 'string', 'bool'
        ]
    
    # 添加代码文本，保留换行符并应用语法高亮
    lines = code.split('\n')
    for i, line in enumerate(lines):
        if i > 0:  # 不是第一行，先添加换行符
            code_para.add_run('\n')
        
        # 语法高亮处理
        # 处理字符串高亮
        string_pattern = r'(\".*?\"|\'.*?\')'
        # 处理注释高亮 - 根据语言调整注释模式
        comment_pattern = r'(//.*?$|#.*?$|/\*.*?\*/|--.*?$|\'\'\'.*?\'\'\'|""".*?""")'
        
        # 先检查是否是注释行
        comment_match = re.search(comment_pattern, line, re.DOTALL)
        if comment_match and comment_match.start() == 0:
            # 整行是注释
            comment_run = code_para.add_run(line)
            comment_run.font.name = 'Consolas'
            comment_run.font.size = Pt(10)
            comment_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色注释
            continue
        
        # 处理字符串和其他部分
        current_pos = 0
        for match in re.finditer(string_pattern + '|' + comment_pattern, line):
            # 添加匹配前的文本
            if match.start() > current_pos:
                pre_text = line[current_pos:match.start()]
                # 检查关键字
                words = re.split(r'(\W+)', pre_text)
                for word in words:
                    if word.strip() in keywords:
                        keyword_run = code_para.add_run(word)
                        keyword_run.font.name = 'Consolas'
                        keyword_run.font.size = Pt(10)
                        keyword_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色关键字
                        keyword_run.bold = True
                    else:
                        normal_run = code_para.add_run(word)
                        normal_run.font.name = 'Consolas'
                        normal_run.font.size = Pt(10)
            
            # 添加匹配的文本（字符串或注释）
            matched_text = match.group(0)
            if re.match(string_pattern, matched_text):
                # 字符串
                string_run = code_para.add_run(matched_text)
                string_run.font.name = 'Consolas'
                string_run.font.size = Pt(10)
                string_run.font.color.rgb = RGBColor(163, 21, 21)  # 红色字符串
            else:
                # 注释
                comment_run = code_para.add_run(matched_text)
                comment_run.font.name = 'Consolas'
                comment_run.font.size = Pt(10)
                comment_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色注释
            
            current_pos = match.end()
        
        # 添加剩余的文本
        if current_pos < len(line):
            remaining_text = line[current_pos:]
            words = re.split(r'(\W+)', remaining_text)
            for word in words:
                if word.strip() in keywords:
                    keyword_run = code_para.add_run(word)
                    keyword_run.font.name = 'Consolas'
                    keyword_run.font.size = Pt(10)
                    keyword_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色关键字
                    keyword_run.bold = True
                else:
                    normal_run = code_para.add_run(word)
                    normal_run.font.name = 'Consolas'
                    normal_run.font.size = Pt(10)
    
    # 添加边框
    try:
        # 为代码块添加左右下边框（上边框由语言标识的下边框提供，或者如果没有语言标识则添加上边框）
        p = code_para._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        border_positions = ['left', 'bottom', 'right']
        if not display_language:
            border_positions.insert(0, 'top')  # 如果没有语言标识，添加上边框
            
        for border_pos in border_positions:
            border = OxmlElement(f'w:{border_pos}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC')
            pBdr.append(border)
        
        pPr.append(pBdr)
        print("Successfully added border to code paragraph")
    except Exception as e:
        print(f"Error adding border to code paragraph: {e}")
    
    # 添加一些空间在代码块后面
    doc.add_paragraph()
    
    return doc


def check_and_process_code_blocks(md_content, doc):
    """直接从Markdown内容中提取和处理代码块，不依赖于HTML转换。"""
    print("直接从Markdown内容中提取和处理代码块...")
    
    # 使用精确的正则表达式匹配代码块
    code_blocks = re.findall(r'```([^\n]*)\n([\s\S]*?)\n```', md_content, re.DOTALL)
    
    if not code_blocks:
        print("未找到代码块")
        return False
    
    print(f"找到 {len(code_blocks)} 个代码块")
    
    # 处理每个代码块
    for i, (lang, code) in enumerate(code_blocks):
        lang = lang.strip()
        code = code.strip()
        
        print(f"处理代码块 {i+1}: 语言='{lang}', 长度={len(code)} 字符")
        
        # 跳过mermaid代码块，因为它们会被单独处理
        if "mermaid" not in lang.lower():
            doc.add_paragraph(f"代码块 {i+1}:", style='Heading 3')
            format_code_block(doc, code, lang)
    
    return True


def debug_code_blocks(md_content):
    """详细分析Markdown内容中的代码块，用于调试"""
    print_section_header("代码块调试信息")
    
    # 1. 使用简单模式查找所有代码块（包括mermaid）
    simple_pattern = r'```(.*?)```'
    simple_blocks = re.findall(simple_pattern, md_content, re.DOTALL)
    debug_print(f"简单模式找到 {len(simple_blocks)} 个代码块", level="INFO")
    
    # 2. 使用精确模式查找所有代码块
    precise_pattern = r'```([^\n]*)\n([\s\S]*?)\n```'
    precise_blocks = re.findall(precise_pattern, md_content, re.DOTALL)
    debug_print(f"精确模式找到 {len(precise_blocks)} 个代码块", level="INFO")
    
    # 3. 分析每个代码块
    for i, (lang, code) in enumerate(precise_blocks):
        lang = lang.strip()
        debug_print(f"代码块 #{i+1}:", level="INFO", indent=1)
        debug_print(f"语言: {lang or '未指定'}", level="INFO", indent=2)
        debug_print(f"长度: {len(code)} 字符", level="INFO", indent=2)
        
        # 检查代码块是否包含特殊字符
        special_chars = [c for c in code if ord(c) > 127]
        if special_chars:
            debug_print(f"包含 {len(special_chars)} 个特殊字符", level="WARNING", indent=2)
            debug_print(f"特殊字符示例: {special_chars[:10]}", level="WARNING", indent=3)
        
        # 检查代码块是否包含可能导致问题的字符
        problem_chars = ['\\', '{', '}', '[', ']', '(', ')', '$', '^', '*', '+', '?', '.', '|']
        found_problem_chars = [c for c in problem_chars if c in code]
        if found_problem_chars:
            debug_print(f"包含可能导致正则表达式问题的字符: {found_problem_chars}", level="WARNING", indent=2)
    
    # 4. 尝试提取mermaid代码块
    mermaid_pattern = r'```mermaid\s+(.*?)\s+```'
    mermaid_blocks = re.findall(mermaid_pattern, md_content, re.DOTALL)
    debug_print(f"找到 {len(mermaid_blocks)} 个mermaid代码块", level="INFO")
    
    # 5. 尝试提取非mermaid代码块
    non_mermaid_pattern = r'```(?!mermaid)([^\n]*)\n([\s\S]*?)\n```'
    non_mermaid_blocks = re.findall(non_mermaid_pattern, md_content, re.DOTALL)
    debug_print(f"找到 {len(non_mermaid_blocks)} 个非mermaid代码块", level="INFO")
    
    print_section_footer("代码块调试信息")
    return precise_blocks


def debug_code_block_processing(md_content, output_file="debug_code_blocks.docx", table_style='Table Grid'):
    """专门用于调试代码块处理的函数"""
    print_section_header("代码块处理调试")
    
    # 创建一个新的Word文档
    doc = Document()
    doc.add_paragraph("代码块处理调试报告", style='Title')
    
    # 1. 添加原始Markdown内容
    doc.add_paragraph("原始Markdown内容", style='Heading 1')
    p = doc.add_paragraph(md_content)
    p.style = 'No Spacing'
    
    # 2. 提取并分析代码块
    doc.add_paragraph("代码块分析", style='Heading 1')
    
    # 使用不同的正则表达式模式提取代码块
    patterns = [
        ("简单模式", r'```(.*?)```'),
        ("精确模式", r'```([^\n]*)\n([\s\S]*?)\n```'),
        ("Mermaid模式", r'```mermaid\s+(.*?)\s+```'),
        ("非Mermaid模式", r'```(?!mermaid)([^\n]*)\n([\s\S]*?)\n```')
    ]
    
    for name, pattern in patterns:
        doc.add_paragraph(f"{name} 提取结果", style='Heading 2')
        debug_print(f"使用 {name} 提取代码块", level="INFO", indent=1)
        
        try:
            if name in ["简单模式", "Mermaid模式"]:
                matches = re.findall(pattern, md_content, re.DOTALL)
                doc.add_paragraph(f"找到 {len(matches)} 个匹配")
                debug_print(f"找到 {len(matches)} 个匹配", level="SUCCESS", indent=2)
                
                for i, match in enumerate(matches):
                    doc.add_paragraph(f"匹配 {i+1}:", style='Heading 3')
                    code_para = doc.add_paragraph(match)
                    code_para.style = 'No Spacing'
                    add_border_to_paragraph(code_para)
            else:
                matches = re.findall(pattern, md_content, re.DOTALL)
                doc.add_paragraph(f"找到 {len(matches)} 个匹配")
                debug_print(f"找到 {len(matches)} 个匹配", level="SUCCESS", indent=2)
                
                for i, (lang, code) in enumerate(matches):
                    doc.add_paragraph(f"匹配 {i+1}:", style='Heading 3')
                    doc.add_paragraph(f"语言: {lang}")
                    code_para = doc.add_paragraph(code)
                    code_para.style = 'No Spacing'
                    add_border_to_paragraph(code_para)
        except Exception as e:
            doc.add_paragraph(f"提取过程出错: {str(e)}")
            debug_print(f"提取过程出错: {str(e)}", level="ERROR", indent=2)
    
    # 3. 测试代码块格式化
    doc.add_paragraph("代码块格式化测试", style='Heading 1')
    debug_print("测试代码块格式化", level="INFO")
    
    # 提取所有代码块
    try:
        precise_blocks = re.findall(r'```([^\n]*)\n([\s\S]*?)\n```', md_content, re.DOTALL)
        debug_print(f"提取到 {len(precise_blocks)} 个代码块进行格式化测试", level="INFO", indent=1)
        
        for i, (lang, code) in enumerate(precise_blocks):
            doc.add_paragraph(f"代码块 {i+1} 格式化:", style='Heading 2')
            debug_print(f"格式化代码块 #{i+1} (语言: {lang.strip() or '未指定'})", level="INFO", indent=2)
            format_code_block(doc, code, lang.strip())
    except Exception as e:
        doc.add_paragraph(f"格式化过程出错: {str(e)}")
        debug_print(f"格式化过程出错: {str(e)}", level="ERROR", indent=1)
    
    # 保存文档
    doc.save(output_file)
    debug_print(f"调试报告已保存到: {output_file}", level="SUCCESS")
    print_section_footer("代码块处理调试")
    
    return output_file


def preprocess_markdown(md_content):
    """预处理Markdown内容，确保表格前后有空行"""
    # 使用正则表达式查找表格并确保其前后有空行
    # 表格通常以 | 开头的行为标志

    
    lines = md_content.split('\n')
    processed_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        processed_lines.append(line)
        
        # 检测表格开始
        if line.strip().startswith('|') and i + 1 < len(lines) and '|' in lines[i+1] and '-' in lines[i+1]:
            # 如果前一行不是空行，添加一个空行
            if i > 0 and processed_lines[-2].strip() != '':
                processed_lines.insert(-1, '')
            
            # 添加表格行直到表格结束
            table_end_idx = i
            while table_end_idx + 1 < len(lines) and lines[table_end_idx + 1].strip().startswith('|'):
                table_end_idx += 1
                processed_lines.append(lines[table_end_idx])
                
            # 如果表格后没有空行，添加一个空行
            if table_end_idx + 1 < len(lines) and lines[table_end_idx + 1].strip() != '':
                processed_lines.append('')
                
            i = table_end_idx
        
        i += 1
    
    return '\n'.join(processed_lines)


def md_to_docx(md_content, output_file=None, debug_mode=False, table_style='Table Grid', mermaid_theme='default', template_file=None):
    """Convert Markdown content to a DOCX file, rendering Mermaid diagrams as images.
    
    Args:
        md_content: Markdown content to convert
        output_file: Optional output file path, defaults to 'output.docx'
        debug_mode: Whether to enable debug mode
        table_style: Style to apply to tables
        mermaid_theme: Theme for Mermaid diagrams ('default', 'dark', 'forest', 'ocean', 'elegant')
        template_file: Optional template docx file to use as base document
        
    Returns:
        The path to the saved DOCX file
    """
    start_time = time.time()
    
    # 处理原始字符串中的转义序列
    if '\\n' in md_content:
        debug_print("检测到转义换行符 (\\n)，进行替换", level="INFO")
        md_content = md_content.replace('\\n', '\n')
    
    if '\\t' in md_content:
        md_content = md_content.replace('\\t', '\t')
    
    if '\\r' in md_content:
        md_content = md_content.replace('\\r', '\r')
    
    # 处理双反斜杠
    if '\\\\' in md_content:
        md_content = md_content.replace('\\\\', '\\')
    
    if debug_mode:
        print_section_header("Markdown转Word转换过程")
        debug_print("启用调试模式", level="INFO")
        debug_code_blocks(md_content)
    
    debug_print(f"处理Markdown内容 (长度: {len(md_content)} 字符)", level="INFO")
    
    # 检查是否有代码块 - 使用更精确的模式
    code_block_pattern = r'```([^\n]*)\n([\s\S]*?)\n```'
    all_code_blocks = re.findall(code_block_pattern, md_content, re.DOTALL)
    code_blocks = []
    debug_print(f"初步检查发现 {len(all_code_blocks)} 个代码块", level="INFO")
    
    for i, (lang, code) in enumerate(all_code_blocks):
        debug_print(f"代码块 #{i+1}: 语言='{lang.strip() or '未指定'}', 长度={len(code)} 字符", level="INFO", indent=1)
        if lang.strip() != "mermaid":
            code_blocks.append((lang.strip(), code.strip()))
    
    # Extract Mermaid blocks
    mermaid_blocks = extract_mermaid_blocks(md_content)
    debug_print(f"发现 {len(mermaid_blocks)} 个Mermaid图表", level="INFO")
    
    # Create a new Word document
    if template_file and os.path.exists(template_file):
        debug_print(f"使用模板文件创建Word文档: {template_file}", level="INFO")
        doc = Document(template_file)
    else:
        debug_print("创建新的Word文档", level="INFO")
        doc = Document()
    
    md_content = preprocess_markdown(md_content)
    debug_print("预处理Markdown内容完成", level="INFO")

    # 创建一个副本用于替换操作
    md_content_copy = md_content
    
    # Replace Mermaid blocks with placeholders and keep track of them
    mermaid_placeholders = []
    if mermaid_blocks:
        debug_print("处理Mermaid图表...", level="INFO")
        
    for i, block in enumerate(mermaid_blocks):
        placeholder = f"MERMAID_DIAGRAM_{i}"
        mermaid_placeholders.append(placeholder)
        # Use regex to replace the block to handle different whitespace patterns
        pattern = r'```mermaid\s+' + re.escape(block) + r'\s+```'
        md_content_copy = re.sub(pattern, placeholder, md_content_copy, flags=re.DOTALL)
        debug_print(f"Mermaid图表 #{i+1} 替换为占位符: {placeholder}", level="INFO", indent=1)
    
    # Replace code blocks with placeholders and keep track of them
    code_placeholders = []
    if code_blocks:
        debug_print("处理代码块...", level="INFO")
        
    for i, (language, block) in enumerate(code_blocks):
        placeholder = f"CODE_BLOCK_{i}"
        code_placeholders.append((placeholder, language, block))
        
        # 使用更精确的模式来替换代码块，保留换行
        if language:
            # 有语言标识的代码块
            pattern = r'```' + re.escape(language) + r'\n' + re.escape(block) + r'\n```'
        else:
            # 无语言标识的代码块
            pattern = r'```\n' + re.escape(block) + r'\n```'
        
        # 尝试替换
        before_replace = md_content_copy
        md_content_copy = re.sub(pattern, placeholder, md_content_copy, count=1, flags=re.DOTALL)
        
        # 如果第一次替换失败，尝试使用更宽松的模式
        if md_content_copy == before_replace:
            debug_print(f"代码块 #{i+1} 第一次替换失败，尝试使用更宽松的模式...", level="WARNING", indent=1)
            if language:
                pattern = r'```' + re.escape(language) + r'[\s\S]*?' + re.escape(block) + r'[\s\S]*?```'
            else:
                pattern = r'```[\s\S]*?' + re.escape(block) + r'[\s\S]*?```'
            
            md_content_copy = re.sub(pattern, placeholder, md_content_copy, count=1, flags=re.DOTALL)
        
        # 检查是否成功替换
        if md_content_copy == before_replace:
            debug_print(f"警告: 代码块 #{i+1} (语言: '{language or '未指定'}') 替换失败", level="WARNING", indent=1)
            # 如果替换失败，尝试直接使用代码内容作为标识符
            md_content_copy = md_content_copy.replace(block, placeholder)
            if md_content_copy == before_replace:
                debug_print(f"警告: 代码块 #{i+1} 直接内容替换也失败", level="WARNING", indent=2)
            else:
                debug_print(f"成功: 代码块 #{i+1} 使用直接内容替换成功", level="SUCCESS", indent=2)
        else:
            debug_print(f"成功: 代码块 #{i+1} 使用模式替换成功", level="SUCCESS", indent=1)
    
    # Combine all placeholders
    all_placeholders = mermaid_placeholders + [p[0] for p in code_placeholders]
    debug_print(f"总占位符数量: {len(all_placeholders)}", level="INFO")
    
    # 如果没有找到任何代码块，但原始内容中确实有代码块标记，则尝试直接处理
    if len(code_blocks) == 0 and '```' in md_content:
        debug_print("未提取到代码块但发现代码块标记，尝试直接处理...", level="WARNING")
        # 直接在文档中添加一个示例代码块
        doc.add_paragraph("以下是从Markdown中提取的代码块：", style='Heading 2')
        
        # 尝试直接匹配和处理代码块
        direct_code_blocks = re.findall(r'```([^\n]*)\n([\s\S]*?)\n```', md_content, re.DOTALL)
        debug_print(f"直接处理模式找到 {len(direct_code_blocks)} 个代码块", level="INFO", indent=1)
        
        for i, (lang, code) in enumerate(direct_code_blocks):
            lang = lang.strip()
            # 保留代码中的换行符，不要去除
            code = code.rstrip()  # 只去除尾部空白，保留换行
            debug_print(f"直接处理 - 代码块 #{i+1}: 语言='{lang or '未指定'}', 长度={len(code)} 字符", level="INFO", indent=2)
            if "mermaid" not in lang.lower():  # 跳过mermaid代码块
                format_code_block(doc, code, lang)
    
    # Convert Markdown to HTML with extensions
    debug_print("将Markdown转换为HTML...", level="INFO")
    html_content = markdown.markdown(
        md_content_copy,
        extensions=[
            'markdown.extensions.extra',
            'markdown.extensions.codehilite',
            'markdown.extensions.tables',
            'markdown.extensions.toc'
        ]
    )
    debug_print(f"HTML转换完成，长度: {len(html_content)} 字符", level="INFO", indent=1)
    
    # Split HTML by placeholders
    parts = []
    if all_placeholders:
        debug_print("按占位符拆分HTML内容...", level="INFO")
        for part in re.split(f"({'|'.join(all_placeholders)})", html_content):
            if part in mermaid_placeholders:
                # This is a mermaid placeholder
                idx = mermaid_placeholders.index(part)
                parts.append(('mermaid', idx))
                debug_print(f"找到Mermaid占位符: {part} (索引 {idx})", level="INFO", indent=1)
            elif any(part == p[0] for p in code_placeholders):
                # This is a code block placeholder
                for i, (placeholder, lang, code) in enumerate(code_placeholders):
                    if part == placeholder:
                        parts.append(('code', i))
                        debug_print(f"找到代码块占位符: {part} (索引 {i})", level="INFO", indent=1)
                        break
            else:
                # This is regular HTML content
                parts.append(('html', part))
                if part.strip():
                    preview = part[:50] + "..." if len(part) > 50 else part
                    debug_print(f"找到HTML内容: {preview}", level="INFO", indent=1)
    else:
        # 如果没有占位符，则将整个内容作为HTML处理
        parts.append(('html', html_content))
        debug_print("未找到占位符，将整个内容作为HTML处理", level="INFO")
    
    debug_print(f"拆分后总部分数: {len(parts)}", level="INFO")
    
    # Process each part
    part_count = {'mermaid': 0, 'code': 0, 'html': 0}
    debug_print("处理各部分内容...", level="INFO")
    
    total_parts = len(parts)
    for idx, (part_type, content) in enumerate(parts):
        part_count[part_type] += 1
        
        # 显示进度
        print_progress(idx + 1, total_parts, prefix="处理进度:", suffix=f"当前: {part_type}")
        
        if part_type == 'mermaid':
            # Render the Mermaid diagram
            mermaid_code = mermaid_blocks[content]
            debug_print(f"渲染Mermaid图表 #{content+1}...", level="INFO", indent=1)
            img_path = render_mermaid_to_image(mermaid_code, theme=mermaid_theme)
            
            if img_path:
                # Add the image to the document
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置图片居中对齐
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(3.5))  # 进一步缩小图片宽度
                debug_print(f"Mermaid图表 #{content+1} 渲染成功并添加到文档", level="SUCCESS", indent=2)
                
                # Clean up the temporary image file
                try:
                    os.unlink(img_path)
                except:
                    pass
            else:
                # If rendering failed, add the Mermaid code as text
                doc.add_paragraph("Failed to render Mermaid diagram:", style='Intense Quote')
                code_para = doc.add_paragraph(mermaid_code)
                code_para.style = 'No Spacing'
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                debug_print(f"Mermaid图表 #{content+1} 渲染失败，作为文本添加", level="ERROR", indent=2)
        
        elif part_type == 'code':
            # Format the code block
            debug_print(f"处理代码块 #{content+1}", level="INFO", indent=1)
            language, code = code_placeholders[content][1], code_placeholders[content][2]
            debug_print(f"代码块语言: '{language or '未指定'}', 长度: {len(code)} 字符", level="INFO", indent=2)
            # 确保代码中的换行被保留
            format_code_block(doc, code, language)
            debug_print(f"代码块 #{content+1} 格式化完成", level="SUCCESS", indent=2)
            
        elif part_type == 'html':
            # Add regular content as paragraphs with proper formatting
            if content.strip():
                html_to_docx(content, doc, table_style)
    
    debug_print(f"处理完成的部分统计: {part_count}", level="INFO")
    
    # Determine output file name if not provided
    if not output_file:
        output_file = 'output.docx'
    
    # Save the document
    debug_print(f"保存文档到: {output_file}", level="INFO")
    doc.save(output_file)
    debug_print(f"文档保存成功", level="SUCCESS")
    
    # 如果没有处理任何代码块，但原始内容中有代码块标记，则创建一个新文档直接处理
    if part_count['code'] == 0 and '```' in md_content:
        debug_print("主流程中未处理任何代码块，创建备份文档进行直接处理", level="WARNING")
        backup_file = output_file.replace('.docx', '_with_code.docx') if output_file.endswith('.docx') else f"{output_file}_with_code.docx"
        
        # 创建一个新文档
        backup_doc = Document()
        backup_doc.add_paragraph("代码块直接处理版本", style='Title')
        
        # 直接处理代码块
        debug_print("在备份文档中直接处理代码块", level="INFO", indent=1)
        check_and_process_code_blocks(md_content, backup_doc)
        
        # 保存备份文档
        backup_doc.save(backup_file)
        debug_print(f"备份文档已保存为: {backup_file}", level="SUCCESS", indent=1)
        
        # 在原始文档中也尝试直接处理代码块
        debug_print("在原始文档中也尝试直接处理代码块", level="INFO", indent=1)
        doc.add_paragraph("代码块", style='Heading 1')
        check_and_process_code_blocks(md_content, doc)
        
        # 重新保存原始文档
        doc.save(output_file)
        debug_print(f"更新后的原始文档已保存", level="SUCCESS", indent=1)
    
    # 计算总耗时
    elapsed_time = time.time() - start_time
    debug_print(f"总耗时: {elapsed_time:.2f} 秒", level="INFO")
    
    if debug_mode:
        print_section_footer("Markdown转Word转换过程")
    
    return output_file 