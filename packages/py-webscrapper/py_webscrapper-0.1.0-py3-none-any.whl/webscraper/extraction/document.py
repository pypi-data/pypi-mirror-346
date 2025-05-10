"""
Document extraction utilities for WebScraper

This module provides functions to extract text from various document formats:
- PDF
- DOCX
- CSV
- Excel
"""

import os
import re
import csv
import logging
import chardet

logger = logging.getLogger(__name__)

def detect_encoding(filepath):
    """
    Detect the encoding of a file
    
    Args:
        filepath: Path to the file
        
    Returns:
        Detected encoding
    """
    with open(filepath, 'rb') as file:
        result = chardet.detect(file.read())
    return result['encoding'] or 'utf-8'

def extract_pdf_text(filepath, logger=None):
    """
    Extract text from PDF file with error handling
    
    Args:
        filepath: Path to PDF file
        logger: Logger instance
        
    Returns:
        Extracted text
    """
    if logger is None:
        logger = logging.getLogger(__name__)
        
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF extraction. Install using: pip install PyPDF2")
        
    text = ""
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Check if PDF is encrypted
            if reader.is_encrypted:
                logger.warning(f"PDF is encrypted, attempting to decrypt: {filepath}")
                try:
                    # Try with empty password
                    reader.decrypt('')
                except:
                    logger.error(f"Failed to decrypt PDF: {filepath}")
                    return "PDF is encrypted and could not be decrypted."
            
            # Check page count
            page_count = len(reader.pages)
            if page_count > 100:
                logger.warning(f"PDF has {page_count} pages, extraction may be slow: {filepath}")
            
            # Extract text from each page
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
                    
                    # Log progress for large documents
                    if page_count > 20 and page_num % 10 == 0:
                        logger.info(f"Extracted {page_num}/{page_count} pages from {filepath}")
                        
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
            
    except Exception as e:
        logger.error(f"Error processing PDF file {filepath}: {str(e)}")
        raise
        
    return text.strip()

def extract_docx_text(filepath, logger=None):
    """
    Extract text from DOCX file with error handling
    
    Args:
        filepath: Path to DOCX file
        logger: Logger instance
        
    Returns:
        Extracted text
    """
    if logger is None:
        logger = logging.getLogger(__name__)
        
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. Install using: pip install python-docx")
        
    try:
        doc = docx.Document(filepath)
        text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text)
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n".join(text)
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {filepath}: {str(e)}")
        raise

def extract_csv_text(filepath, logger=None):
    """
    Extract text from CSV file with error handling
    
    Args:
        filepath: Path to CSV file
        logger: Logger instance
        
    Returns:
        Extracted text
    """
    if logger is None:
        logger = logging.getLogger(__name__)
        
    text = []
    
    try:
        # Detect encoding 
        encoding = detect_encoding(filepath)
        
        # Try multiple delimiters
        delimiters = [',', ';', '\t', '|']
        best_delimiter = ','
        max_columns = 0
        
        # Find the best delimiter by checking which one gives the most columns
        for delimiter in delimiters:
            try:
                with open(filepath, 'r', encoding=encoding) as csvfile:
                    sample = csvfile.read(2048)  # Read a sample
                    if delimiter in sample:
                        sample_rows = sample.split('\n')
                        if len(sample_rows) > 1:
                            cols = len(sample_rows[0].split(delimiter))
                            if cols > max_columns:
                                max_columns = cols
                                best_delimiter = delimiter
            except Exception:
                continue
        
        # Read with best delimiter
        with open(filepath, 'r', encoding=encoding) as csvfile:
            csv_reader = csv.reader(csvfile, delimiter=best_delimiter)
            for row in csv_reader:
                text.append(" ".join([str(cell) for cell in row if cell]))
                
        return "\n".join(text)
        
    except Exception as e:
        logger.error(f"Error extracting text from CSV file {filepath}: {str(e)}")
        raise

def extract_excel_text(filepath, logger=None):
    """
    Extract text from Excel file with error handling
    
    Args:
        filepath: Path to Excel file
        logger: Logger instance
        
    Returns:
        Extracted text
    """
    if logger is None:
        logger = logging.getLogger(__name__)
        
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl is required for Excel extraction. Install using: pip install openpyxl")
        
    text = []
    
    try:
        workbook = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        
        # Track progress for large files
        sheet_count = len(workbook.sheetnames)
        if sheet_count > 10:
            logger.info(f"Processing large Excel file with {sheet_count} sheets: {filepath}")
        
        for sheet_name in workbook.sheetnames:
            try:
                sheet = workbook[sheet_name]
                
                # Add sheet name as header
                text.append(f"Sheet: {sheet_name}")
                
                # Process rows
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:  # Skip empty rows
                        text.append(" | ".join(row_values))
                    row_count += 1
                    
                    # Log progress for large sheets
                    if row_count > 1000 and row_count % 1000 == 0:
                        logger.info(f"Processed {row_count} rows in sheet {sheet_name}")
                        
                text.append("")  # Add separator between sheets
            except Exception as e:
                logger.warning(f"Error processing sheet {sheet_name}: {str(e)}")
                continue
        
        return "\n".join(text)
        
    except Exception as e:
        logger.error(f"Error extracting text from Excel file {filepath}: {str(e)}")
        raise 