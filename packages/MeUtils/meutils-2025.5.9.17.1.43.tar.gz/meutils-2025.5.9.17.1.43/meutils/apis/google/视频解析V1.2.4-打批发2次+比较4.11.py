import re
import requests
import json
import os
import shutil  # 移到顶部导入
from urllib.parse import urlparse
from docx import Document
import glob
import datetime

# 配置信息
config = {
    "api_url": "https://api.chatfire.cn/v1/chat/completions",
    "api_key": "sk-",
    "input_txt": "input/video_urls.txt",  # 包含视频 URL 的输入文件
}

def get_timestamped_output_dir():
    # 生成带时间戳的输出目录
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join('output', timestamp)
    
    # 确保主目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 创建子目录 A, B, C
    for subdir in ['A', 'B', 'C']:
        subdir_path = os.path.join(output_dir, subdir)
        os.makedirs(subdir_path, exist_ok=True)
    
    return output_dir

def analyze_video(video_url, prompt):
    """使用API分析视频"""
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {config["api_key"]}'
    }
    payload = {
        "model": "gemini-all",
        "messages": [
            {
                "role": "system",
                "content": """你是一位专业的短剧编剧。请将视频内容转换为剧本格式，重点关注：
                1.  角色对白（包括语气、情感提示）
                2. 具体的表演动作指示
                3. 场景和道具说明
                4. 表情和肢体语言
                5. 非台词的字幕
                6.人物台词不要出现张冠李戴的情况，无法判别人物身份时，这一集中出现最多的年轻成年女性为女主，出现最多的年轻成年男性为男主
                7. 对重复的台词、动作进行识别和删除
                8.每一集的场次都需要从1开始，用1.2.3连续符号标注，不要被上一集影响
                9. 剧本全文不要出现"*"
                10. 注意分辨不同人物，学习我所给出的主要人物关系"""
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": prompt
                    },
                    {
                        "type": "video_url",
                        "video_url": {
                            "url": video_url
                        }
                    }
                ]
            }
        ]
    }
    try:
        print(f"正在分析视频: {video_url}")
        response = requests.post(
            config["api_url"],
            headers=headers,
            json=payload,
            timeout=180
        )
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('message', {}).get('content')
        else:
            print(f"错误信息：{response.text}")
            return None
    except Exception as e:
        print(f"分析失败: {e}")
        return None


def format_episode_title(filename):
    # 从文件名中去掉.txt后缀，前面加上"第"，后面加上"集"
    name = os.path.splitext(filename)[0]
    return f"第{name}集"

def process_scene(text):
    # 提取场次信息和其余内容
    scene_match = re.match(r'(场次\s*[\[【]?\s*\d+\s*[\]】]?.*?)(?=\n)', text, re.DOTALL)
    if scene_match:
        scene_header = scene_match.group(1).strip()
        remaining_content = text[scene_match.end():].strip()
        return scene_header, remaining_content
    return None, text

def filter_content(text):
    # 定义需要过滤的废话
    filters = [
        "剧本结束。",
        "好的，请看剧本：",
        "请看剧本：",
        "以下是剧本：",
        "剧本如下：",
    ]# 过滤掉这些废话
    filtered_text = text
    for filter_text in filters:
        filtered_text = filtered_text.replace(filter_text, "")
    return filtered_text.strip()

def merge_to_word(txt_files, output_file):  # 修改参数
    document = Document()
    for txt_file in txt_files:  # 直接使用文件列表
        episode_title = format_episode_title(os.path.basename(txt_file))  # 使用basename
        document.add_heading(episode_title, level=2)
        with open(txt_file, 'r', encoding='utf-8') as f:  # 直接使用文件路径
            content = f.read()        
        content = filter_content(content)
        scenes = [s for s in content.split('\n\n') if s.strip()]
        
        for scene in scenes:
            scene_header, scene_content = process_scene(scene)
            if scene_header:
                document.add_heading(scene_header, level=3)
                if scene_content:
                    document.add_paragraph(scene_content)
            else:
                document.add_paragraph(scene)

    document.save(output_file)
    print(f"文档已保存至: {output_file}")


    
def main():
    try:
        # 创建带时间戳的输出目录
        output_dir = get_timestamped_output_dir()
        
        # 确保所有子目录都已创建
        for subdir in ['A', 'B', 'C']:
            subdir_path = os.path.join(output_dir, subdir)
            if not os.path.exists(subdir_path):
                os.makedirs(subdir_path, exist_ok=True)
        
        # 读取视频URL列表
        with open(config['input_txt'], 'r') as f:
            urls = [line.strip() for line in f if line.strip()]        
        # 通用的剧本生成prompt（保持不变）
        base_prompt = """出场人物

场次编号+地点+时间（日/夜）+内/外景
（场景描述包含布景要求、核心道具、环境氛围）
字幕：角色身份说明（角色首次出场时标注）
旁白：画外音内容（如有）
△ 动作描述（使用△符号起始，包含肢体语言、关键走位）
角色名（情感提示）：对白内容（保留方言特色如"啷个""铲铲"等）

【剧本示例参考】
出场人物：
陈海清：冷冻机厂副主任
刘厂长：冷冻机三分厂厂长
蔡晓艳：冷冻机一分厂车间组长
王日新：冷冻机二分厂财务主任
李医生：医生
赵春红：蔡晓艳嫂子
蔡晓艳母亲
蔡晓艳儿子：邵一帆

1. 川南冷冻厂 夜 外
字幕：1981年
长江中上游洪水灾后
旁白：长江中上游洪水灾后，群众正在积极投入恢复建设中……
△  灾难现场，一双男人的手把掉在地上的招牌捡起来，招牌上写着“川南冷冻机械二分厂”。
工人甲：快来帮忙！
工人乙：来了来了！
△  一名工人拉木头时不小心脚滑了一下，拖车失控，一只戴了手表的男的人手突然握住把手，把车控制住。
字幕：冷冻机总厂副主任陈海清陈海清：我来帮你。
△  一群厂里的管理员赶忙追过来。
管理员（劝阻）：领导领导，陈主任……
陈海清（忙着干活）：拿个石头——
管理员：我们分厂的这个灾后工作做得很好，领导，你不用亲自来的嘛？
陈海清：莫浪费时间，救灾要紧！
△  突然从厂区二楼医务室传来一个女人的吵闹声。女人（画外音）：领导算个屁，李医生你再不跟老子走，老子把你医务室给你砸烂，你个背时（倒霉）领导，脸皮比城墙加拐棍还厚！
△  陈海清跟现场工人都抬头看向医务室。

2. 医务室夜 内
字幕：蔡晓艳冷东机二分厂车间组长
蔡晓艳：李医生，我嫂子怀起娃儿（怀着孕），困到车间头好几天了，现在大出血！工人里头好多女工，骨头都砸断了，再不去真的来不及了！走嘛！
△  陈海清走到医务室门口。李医生：但是我这儿，还没检查完的嘛。
......

依照场景类推

【处理要求】
1. 严格区分日/夜、内/外景，每场戏用编号分隔
2. 人物首次出场必须用字幕说明姓名+身份
3. 动作描述用△符号，包含表情细节（如"揩了揩泪"）
4. 保留方言词汇和口语化表达（如"歪婆娘""背时"）
5. 注意角色关系标注（如"蔡晓艳儿子邵一帆"）
6. 关键道具需特别说明（如"二八大杠""蜂窝煤"）
7. 情感提示用括号标注（如"憋着气""得意一笑"）
8. 人物台词前面不要加△ ，只有人物动作才加△ 
9. 场景描述不需要加上“场景描述：”和括号
10. 对生成内容进行检查，对重复内容进行删除（重复的对话、场景描述）
11. 禁止人物台词张冠李戴

请确保生成格式与示例完全一致，包含所有叙事元素：场景、字幕、旁白、动作、对白、情感提示，剧本全文不要出现“*”。并对重复内容进行检查和删除。

并注意：
1.不需要任何额外的解释、问候或结束语（如，
好的，请看剧本：
好的我给你生成剧本
剧本结束了希望你满意
）诸如此类都不要给！
2. 直接输出剧本内容
3. 保持格式规范和一致性

"""
        
        # 存储所有生成的txt文件路径，按子目录分类        
        txt_files = {'A': [], 'B': [], 'C': []}
        
        # 处理每个视频URL
        for index, url in enumerate(urls, start=1):
            # 提取文件名作为输出文件名
            parsed_url = urlparse(url)
            filename = os.path.basename(parsed_url.path)
            
            # 为子目录A和B分别生成剧本文件路径
            output_filename_A = os.path.join(output_dir, 'A', f"{index}.txt")
            output_filename_B = os.path.join(output_dir, 'B', f"{index}.txt")
            output_filename_C = os.path.join(output_dir, 'C', f"{index}.txt")
            
            # 为每个URL生成两次剧本
            print(f"为第{index}集生成两份剧本...")
            
            # 生成第一个剧本 (A)
            analysis_result_A = analyze_video(url, base_prompt)
            success_A = False
            if analysis_result_A:
                # 保存剧本到txt文件
                with open(output_filename_A, 'w', encoding='utf-8') as output_file:
                    output_file.write(analysis_result_A)
                print(f"第{index}集剧本A已保存到 {output_filename_A}")
                txt_files['A'].append(output_filename_A)
                success_A = True
            else:
                print(f"第{index}集视频分析A失败")
            
            # 生成第二个剧本 (B)
            analysis_result_B = analyze_video(url, base_prompt)
            success_B = False
            if analysis_result_B:
                # 保存剧本到txt文件
                with open(output_filename_B, 'w', encoding='utf-8') as output_file:
                    output_file.write(analysis_result_B)
                print(f"第{index}集剧本B已保存到 {output_filename_B}")
                txt_files['B'].append(output_filename_B)
                success_B = True
            else:
                print(f"第{index}集视频分析B失败")
            
            # 处理生成结果 - 这部分应该在每个视频的循环内部
            # 1. 一个成功一个失败的情况
            if success_A and not success_B:
                # 确保目录存在
                os.makedirs(os.path.dirname(output_filename_C), exist_ok=True)
                # 复制文件
                shutil.copy(output_filename_A, output_filename_C)
                txt_files['C'].append(output_filename_C)
                print(f"第{index}集剧本A成功而B失败，已将A复制到C目录")
            elif not success_A and success_B:
                os.makedirs(os.path.dirname(output_filename_C), exist_ok=True)
                shutil.copy(output_filename_B, output_filename_C)
                txt_files['C'].append(output_filename_C)
                print(f"第{index}集剧本B成功而A失败，已将B复制到C目录")
            # 2. 检查文件大小异常
            elif success_A and success_B:
                # 获取文件大小
                size_A = os.path.getsize(output_filename_A)
                size_B = os.path.getsize(output_filename_B)
                
                # 检查是否有异常大的文件（超过6KB）
                max_normal_size = 6 * 1024  # 6KB
                
                if size_A > max_normal_size and size_B <= max_normal_size:
                    os.makedirs(os.path.dirname(output_filename_C), exist_ok=True)
                    shutil.copy(output_filename_B, output_filename_C)
                    txt_files['C'].append(output_filename_C)
                    print(f"第{index}集剧本A文件异常大({size_A}字节)，已将B复制到C目录")
                elif size_B > max_normal_size and size_A <= max_normal_size:
                    os.makedirs(os.path.dirname(output_filename_C), exist_ok=True)
                    shutil.copy(output_filename_A, output_filename_C)
                    txt_files['C'].append(output_filename_C)
                    print(f"第{index}集剧本B文件异常大({size_B}字节)，已将A复制到C目录")
                # 3. 两个都正常，比较哪个更好
                else:
                    # 这里简单比较文本长度，更复杂的比较方法可以根据需要添加
                    if len(analysis_result_A) > len(analysis_result_B):
                        os.makedirs(os.path.dirname(output_filename_C), exist_ok=True)
                        shutil.copy(output_filename_A, output_filename_C)
                        txt_files['C'].append(output_filename_C)
                        print(f"第{index}集剧本A和B都成功，A内容更丰富，已将A复制到C目录")
                    else:
                        os.makedirs(os.path.dirname(output_filename_C), exist_ok=True)
                        shutil.copy(output_filename_B, output_filename_C)
                        txt_files['C'].append(output_filename_C)
                        print(f"第{index}集剧本A和B都成功，B内容更丰富，已将B复制到C目录")
        
        # 根据子目录C的结果生成Word文档
        if txt_files['C']:
            output_docx = os.path.join(output_dir, 'merged_scripts.docx')
            merge_to_word(txt_files['C'], output_docx)
            print(f"剧本已合并到 {output_docx}")
        else:
            print("没有成功生成任何剧本文件")
    except Exception as e:
        print(f"发生错误：{str(e)}")
        
# 运行主函数
if __name__ == "__main__":
    main()
